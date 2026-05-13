from __future__ import annotations

import copy
import re
import shutil
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from lxml import etree


SRC = Path(r"C:\Users\hy971\source\repos\MyDietitianMobileApp\tmp-build\new_son_Dytopia_Bitirme_Tezi_Final_v7.docx")
OUT = Path(r"C:\Users\hy971\source\repos\MyDietitianMobileApp\tmp-build\Dytopia_Bitirme_Tezi_Final_Referans_Formatli.docx")
WORK = Path(r"C:\Users\hy971\source\repos\MyDietitianMobileApp\tmp-build\_format_work.docx")

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "xml": "http://www.w3.org/XML/1998/namespace",
    "ct": "http://schemas.openxmlformats.org/package/2006/content-types",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
}
W = NS["w"]
R = NS["r"]


def norm(text: str) -> str:
    return " ".join(text.split())


def tr_upper(text: str) -> str:
    return text.replace("i", "İ").replace("ı", "I").upper()


def collect_between(doc: Document, start_heading: str, end_heading: str | None = None) -> list[str]:
    items: list[str] = []
    active = False
    for p in doc.paragraphs:
        txt = norm(p.text)
        if not txt:
            continue
        if p.style.name.startswith("Heading") and txt.upper() == start_heading.upper():
            active = True
            continue
        if active and p.style.name.startswith("Heading"):
            if end_heading is None or txt.upper() == end_heading.upper():
                break
        if active:
            items.append(txt)
    return items


def extract_metadata(doc: Document) -> dict[str, object]:
    paragraphs = [norm(p.text) for p in doc.paragraphs]
    title = next((t for t in paragraphs if len(t) > 40 and "TARİF" in t.upper()), paragraphs[5])
    name = next((t for t in paragraphs if "YILDIRIM" in t.upper()), "Halil İbrahim YILDIRIM")
    student_no = next((re.sub(r"^OKUL\s*NO\s*:\s*", "", t, flags=re.I) for t in paragraphs if "OKUL NO" in t.upper()), "212503009")
    advisor_line = next((t for t in paragraphs if "Danışman:" in t), "Danışman: Dr. Öğr. Üyesi Volkan ATEŞ")
    advisor = advisor_line.split(":", 1)[1].strip() if ":" in advisor_line else advisor_line.strip()

    ozet_raw = collect_between(doc, "ÖZET", "ABSTRACT")
    abstract_raw = collect_between(doc, "ABSTRACT", "ÖNSÖZ")
    onsoz_raw = collect_between(doc, "ÖNSÖZ", "İÇİNDEKİLER")
    abbr_raw = collect_between(doc, "KISALTMALAR", "TABLOLAR LİSTESİ")

    def clean_summary(items: list[str], keyword_label: str) -> tuple[list[str], str]:
        body: list[str] = []
        keywords = ""
        for item in items:
            up = item.upper()
            up_ascii = up.replace("İ", "I")
            if (
                tr_upper(item) == tr_upper(title)
                or "YILDIRIM" in up
                or "OKUL NO" in up
                or "LISANS BITIRME" in up_ascii
                or "GRADUATION THESIS" in up
                or up.startswith("A DIETITIAN RULE")
            ):
                continue
            if up.startswith(keyword_label.upper()):
                keywords = item
            else:
                body.append(item)
        return body, keywords

    ozet, anahtar = clean_summary(ozet_raw, "Anahtar")
    abstract, keywords = clean_summary(abstract_raw, "Keywords")

    abbr: list[tuple[str, str]] = []
    for item in abbr_raw:
        if ":" in item:
            key, val = item.split(":", 1)
            abbr.append((key.strip(), val.strip()))
        else:
            parts = item.split(maxsplit=1)
            if len(parts) == 2:
                abbr.append((parts[0], parts[1]))
    if not abbr:
        abbr = [
            ("API", "Application Programming Interface"),
            ("EF Core", "Entity Framework Core"),
            ("JWT", "JSON Web Token"),
            ("LLM", "Large Language Model"),
            ("MVP", "Minimum Viable Product"),
            ("REST", "Representational State Transfer"),
            ("UI", "User Interface"),
            ("UX", "User Experience"),
        ]

    return {
        "title": title,
        "name": name,
        "student_no": student_no,
        "advisor": advisor,
        "ozet": ozet,
        "anahtar": anahtar,
        "abstract": abstract,
        "keywords": keywords,
        "onsoz": onsoz_raw,
        "abbr": abbr,
    }


def polish_existing_body(path: Path) -> None:
    doc = Document(path)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(0.7)
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        if style_name == "Heading 1":
            style.font.size = Pt(12)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            style.font.size = Pt(12)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    body_started = False
    major = {
        "GİRİŞ",
        "SONUÇ VE ÖNERİLER",
        "KAYNAKÇA",
        "EKLER",
        "ÖZGEÇMİŞ",
    }
    for p in doc.paragraphs:
        txt = norm(p.text)
        if txt == "GİRİŞ":
            body_started = True
        if not body_started:
            continue
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(0, 0, 0)
        if p.style.name == "Normal" and txt:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0.7)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)
        if p.style.name == "Heading 1":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(12)
        if p.style.name.startswith("Heading") and txt in major:
            p.paragraph_format.page_break_before = True

    for table in doc.tables:
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(2)
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(path)


def qn(tag: str) -> str:
    prefix, local = tag.split(":")
    return f"{{{NS[prefix]}}}{local}"


def el(tag: str, attrs: dict[str, str] | None = None, children: list[etree._Element] | None = None) -> etree._Element:
    node = etree.Element(qn(tag))
    for key, value in (attrs or {}).items():
        node.set(qn(key) if ":" in key else key, value)
    for child in children or []:
        node.append(child)
    return node


def text_of_element(p: etree._Element) -> str:
    return "".join(p.xpath(".//w:t/text()", namespaces=NS))


def run_props(size: int = 22, bold: bool = False, italic: bool = False) -> etree._Element:
    rpr = el("w:rPr")
    rpr.append(el("w:rFonts", {"w:ascii": "Times New Roman", "w:hAnsi": "Times New Roman", "w:cs": "Times New Roman"}))
    if bold:
        rpr.append(el("w:b"))
        rpr.append(el("w:bCs"))
    if italic:
        rpr.append(el("w:i"))
        rpr.append(el("w:iCs"))
    rpr.append(el("w:sz", {"w:val": str(size)}))
    rpr.append(el("w:szCs", {"w:val": str(size)}))
    rpr.append(el("w:color", {"w:val": "000000"}))
    return rpr


def paragraph(
    text: str = "",
    *,
    align: str | None = None,
    bold: bool = False,
    italic: bool = False,
    size: int = 22,
    before: int = 0,
    after: int = 120,
    line: int | None = None,
    first: int | None = None,
    left: int | None = None,
    tabs: bool = False,
) -> etree._Element:
    p = el("w:p")
    ppr = el("w:pPr")
    if align:
        ppr.append(el("w:jc", {"w:val": align}))
    spacing_attrs = {"w:before": str(before), "w:after": str(after)}
    if line:
        spacing_attrs.update({"w:line": str(line), "w:lineRule": "auto"})
    ppr.append(el("w:spacing", spacing_attrs))
    ind_attrs: dict[str, str] = {}
    if first is not None:
        ind_attrs["w:firstLine"] = str(first)
    if left is not None:
        ind_attrs["w:left"] = str(left)
    if ind_attrs:
        ppr.append(el("w:ind", ind_attrs))
    if tabs:
        tabs_el = el("w:tabs")
        tabs_el.append(el("w:tab", {"w:val": "right", "w:leader": "dot", "w:pos": "8500"}))
        ppr.append(tabs_el)
    p.append(ppr)
    if text:
        parts = text.split("\t")
        for idx, part in enumerate(parts):
            if idx:
                p.append(el("w:r", children=[el("w:tab")]))
            r = el("w:r")
            r.append(run_props(size=size, bold=bold, italic=italic))
            t = el("w:t", {"xml:space": "preserve"})
            t.text = part
            r.append(t)
            p.append(r)
    return p


def page_break() -> etree._Element:
    return el("w:p", children=[el("w:r", children=[el("w:br", {"w:type": "page"})])])


def sect_pr(*, footer_rid: str | None, fmt: str, start: int, next_page: bool = True) -> etree._Element:
    sp = el("w:sectPr")
    if footer_rid:
        sp.append(el("w:footerReference", {"w:type": "default", "r:id": footer_rid}))
    if next_page:
        sp.append(el("w:type", {"w:val": "nextPage"}))
    sp.append(el("w:pgSz", {"w:w": "11906", "w:h": "16838"}))
    sp.append(el("w:pgMar", {"w:top": "1440", "w:right": "1440", "w:bottom": "1440", "w:left": "1701", "w:header": "708", "w:footer": "708", "w:gutter": "0"}))
    sp.append(el("w:pgNumType", {"w:fmt": fmt, "w:start": str(start)}))
    sp.append(el("w:cols", {"w:space": "708"}))
    sp.append(el("w:docGrid", {"w:linePitch": "360"}))
    return sp


def section_break(*, footer_rid: str | None, fmt: str, start: int) -> etree._Element:
    p = el("w:p")
    ppr = el("w:pPr")
    ppr.append(sect_pr(footer_rid=footer_rid, fmt=fmt, start=start, next_page=True))
    p.append(ppr)
    return p


def add_page(elements: list[etree._Element], page: list[etree._Element], break_after: bool = True) -> None:
    elements.extend(page)
    if break_after:
        elements.append(page_break())


def toc_line(label: str, page: str, *, bold: bool = False, indent: int = 0) -> etree._Element:
    return paragraph(f"{label}\t{page}", bold=bold, size=20, after=60, left=indent, tabs=True)


def front_matter(meta: dict[str, object], roman_footers: list[str]) -> list[etree._Element]:
    title = tr_upper(str(meta["title"]))
    name = str(meta["name"])
    name_up = tr_upper(name)
    student_no = str(meta["student_no"])
    advisor = str(meta["advisor"])
    ozet = list(meta["ozet"]) or ["Bu çalışmada, diyetisyen kontrollü beslenme süreçlerinde danışanın mevcut malzemeleri ile diyetisyen tarafından tanımlanan tarif kuralları arasındaki uyumu açıklanabilir biçimde değerlendiren web ve mobil tabanlı bir akıllı mutfak asistanı geliştirilmiştir."]
    abstract = list(meta["abstract"]) or ["In this study, a web and mobile based smart kitchen assistant was developed to evaluate the compatibility between ingredients available in a client's kitchen and dietitian-defined recipe rules in an explainable manner."]
    onsoz = list(meta["onsoz"]) or ["Bu tez çalışması, bilgisayar mühendisliği bakış açısıyla beslenme takibi, mobil sağlık uygulamaları, kural tabanlı öneri sistemleri ve SaaS mimarisi konularını bir araya getiren uygulamalı bir mezuniyet projesi olarak hazırlanmıştır."]
    abbr = list(meta["abbr"])

    elems: list[etree._Element] = []
    cover_common = [
        paragraph("T.C.", align="center", bold=True, size=24, after=160),
        paragraph("TARSUS ÜNİVERSİTESİ", align="center", bold=True, size=24, after=160),
        paragraph("MÜHENDİSLİK FAKÜLTESİ", align="center", bold=True, size=24, after=160),
        paragraph("BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ", align="center", bold=True, size=24, after=840),
    ]

    add_page(
        elems,
        [
            paragraph("", after=900),
            *cover_common,
            paragraph(title, align="center", bold=True, size=22, after=800),
            paragraph(name_up, align="center", bold=True, size=22, after=180),
            paragraph(f"OKUL NO: {student_no}", align="center", size=22, after=780),
            paragraph("LİSANS BİTİRME TEZİ", align="center", bold=True, size=22, after=900),
            paragraph("TARSUS - 2026", align="center", bold=True, size=22, after=0),
        ],
    )

    add_page(
        elems,
        [
            paragraph("", after=700),
            *copy.deepcopy(cover_common),
            paragraph(title, align="center", bold=True, size=22, after=650),
            paragraph(name_up, align="center", bold=True, size=22, after=150),
            paragraph(f"OKUL NO: {student_no}", align="center", size=22, after=520),
            paragraph("LİSANS BİTİRME TEZİ", align="center", bold=True, size=22, after=360),
            paragraph("Danışman", align="center", bold=True, size=22, after=80),
            paragraph(advisor, align="center", size=22, after=620),
            paragraph("TARSUS - 2026", align="center", bold=True, size=22, after=0),
        ],
    )

    accept_text = (
        f"{name_up} tarafından hazırlanan \"{title}\" başlıklı bu çalışma, "
        "Tarsus Üniversitesi Mühendislik Fakültesi Bilgisayar Mühendisliği Bölümü "
        "Lisans Bitirme Tezi olarak kabul edilmiştir."
    )
    add_page(
        elems,
        [
            paragraph("", after=620),
            paragraph("KABUL VE ONAY SAYFASI", align="center", bold=True, size=22, after=520),
            paragraph(accept_text, size=20, after=360, line=300, first=0),
            paragraph("Tez Teslim Tarihi:  .../.../2026", size=20, after=360, first=0),
            paragraph("Danışman\tİmza", bold=True, size=20, after=160, tabs=True),
            paragraph(f"{advisor}\t____________________", size=20, after=320, tabs=True),
            paragraph("Jüri Üyesi\tİmza", bold=True, size=20, after=160, tabs=True),
            paragraph("Jüri Üyesi Adı Soyadı\t____________________", size=20, after=320, tabs=True),
            paragraph("Jüri Üyesi\tİmza", bold=True, size=20, after=160, tabs=True),
            paragraph("Jüri Üyesi Adı Soyadı\t____________________", size=20, after=1800, tabs=True),
            paragraph("Bu sayfadaki jüri adı ve tarih alanları nihai teslim öncesinde bölüm tarafından bildirilen kesin bilgilerle güncellenecektir.", size=18, after=0, line=280),
        ],
    )

    ethics_1 = (
        "Bu tez çalışmasının hazırlanması ve yazımı sırasında bilimsel araştırma ve yayın etiği kurallarına uyduğumu; "
        "bu çalışma kapsamında elde edilen tüm bilgi, belge, değerlendirme ve sonuçları akademik dürüstlük ilkeleri doğrultusunda sunduğumu; "
        "yararlandığım bütün kaynaklara metin içinde ve kaynakça bölümünde uygun biçimde atıf yaptığımı beyan ederim."
    )
    ethics_2 = (
        "Bu çalışmada kullanılan yazılım çıktıları, test verileri ve prototip ekranları araştırma ve mezuniyet projesi amaçlıdır. "
        "Geliştirilen sistem klinik karar verme ya da profesyonel diyetisyen hizmetinin yerine geçecek bir ürün olarak değil, "
        "malzeme normalizasyonu ve kural tabanlı tarif önerisi problemini inceleyen akademik bir prototip olarak değerlendirilmiştir."
    )
    elems.extend(
        [
            paragraph("", after=620),
            paragraph("ETİK BEYAN", align="center", bold=True, size=22, after=520),
            paragraph(ethics_1, size=20, after=160, line=300, first=420),
            paragraph(ethics_2, size=20, after=520, line=300, first=420),
            paragraph("Tarih:\t.../.../2026", size=20, after=360, tabs=True),
            paragraph("İmza:\t____________________________", size=20, after=360, tabs=True),
            paragraph(f"Adı Soyadı:\t{name_up}", size=20, after=0, tabs=True),
            section_break(footer_rid=None, fmt="decimal", start=1),
        ]
    )

    elems.extend(
        [
            paragraph("", after=760),
            paragraph("ÖZET", align="center", bold=True, size=22, after=520),
            *[paragraph(x, size=20, after=120, line=300, first=420) for x in ozet],
            paragraph(str(meta["anahtar"] or "Anahtar Kelimeler: Akıllı tarif eşleştirme, malzeme normalizasyonu, diyet uyum takibi, kural tabanlı öneri sistemi, mobil sağlık."), size=20, after=0, line=280, first=0),
            section_break(footer_rid=roman_footers[0], fmt="decimal", start=1),
        ]
    )
    elems.extend(
        [
            paragraph("", after=760),
            paragraph("ABSTRACT", align="center", bold=True, size=22, after=520),
            *[paragraph(x, size=20, after=120, line=300, first=420) for x in abstract],
            paragraph(str(meta["keywords"] or "Keywords: Smart recipe matching, ingredient normalization, diet compliance tracking, rule-based recommender system, mobile health."), size=20, after=0, line=280, first=0),
            section_break(footer_rid=roman_footers[1], fmt="decimal", start=1),
        ]
    )
    elems.extend(
        [
            paragraph("", after=760),
            paragraph("ÖNSÖZ", align="center", bold=True, size=22, after=520),
            *[paragraph(x, size=20, after=120, line=300, first=420) for x in onsoz],
            paragraph("", after=420),
            paragraph(name_up, align="right", bold=True, size=20, after=80),
            paragraph("Tarsus, 2026", align="right", size=20, after=0),
            section_break(footer_rid=roman_footers[2], fmt="decimal", start=1),
        ]
    )

    abbr_page = [paragraph("", after=620), paragraph("KISALTMALAR", align="center", bold=True, size=22, after=420)]
    for key, val in abbr:
        abbr_page.append(paragraph(f"{key}\t{val}", size=20, after=36, tabs=False, left=260))
    elems.extend(abbr_page)
    elems.append(section_break(footer_rid=roman_footers[3], fmt="decimal", start=1))

    toc_1 = [
        paragraph("", after=420),
        paragraph("İçindekiler", align="center", bold=True, size=20, after=420),
        toc_line("ÖZET", "vi", bold=True),
        toc_line("ABSTRACT", "vii", bold=True),
        toc_line("ÖNSÖZ", "viii", bold=True),
        toc_line("KISALTMALAR", "ix", bold=True),
        toc_line("İÇİNDEKİLER", "x", bold=True),
        toc_line("TABLOLAR LİSTESİ", "xii", bold=True),
        toc_line("ŞEKİLLER LİSTESİ", "xiii", bold=True),
        toc_line("EKLER LİSTESİ", "xiv", bold=True),
        toc_line("GİRİŞ", "1", bold=True),
        toc_line("Problem Tanımı", "1", indent=360),
        toc_line("Araştırma Sorusu", "2", indent=360),
        toc_line("Amaç ve Kapsam", "2", indent=360),
        toc_line("Çalışmanın Özgün Değeri", "2", indent=360),
        toc_line("İş Paketleri ve Zaman Planı", "3", indent=360),
        toc_line("BİRİNCİ BÖLÜM    TEORİK ALT YAPI", "3", bold=True),
        toc_line("1.1 Mobil Sağlık ve Dijital Beslenme Sistemleri", "3", indent=360),
        toc_line("1.2 Beslenme Öneri Sistemleri", "4", indent=360),
        toc_line("1.3 Gıda Ontolojileri ve Malzeme Standardizasyonu", "4", indent=360),
        toc_line("1.4 Bulanık Eşleştirme ve Normalizasyon Katmanları", "5", indent=360),
        toc_line("1.5 Açıklanabilir Öneri Sistemleri", "5", indent=360),
        toc_line("1.6 Tarif Önerilerinde Bilgi Grafı ve Kural Tabanlı Yaklaşımlar", "5", indent=360),
        toc_line("1.7 Literatür Özeti ve Çalışmanın Konumu", "6", indent=360),
    ]
    elems.extend(toc_1)
    elems.append(section_break(footer_rid=roman_footers[4], fmt="decimal", start=1))

    toc_2 = [
        paragraph("", after=360),
        toc_line("İKİNCİ BÖLÜM    MATERYAL VE METOT", "6", bold=True),
        toc_line("2.1 Kullanılan Yazılım ve Donanım Ortamı", "6", indent=360),
        toc_line("2.2 Veri Modeli", "9", indent=360),
        toc_line("2.3 Malzeme Normalizasyonu", "9", indent=360),
        toc_line("2.4 Akıllı Tarif Şeması", "11", indent=360),
        toc_line("2.5 Kural Tabanlı Karşılaştırma Motoru", "11", indent=360),
        toc_line("2.6 Access Key ile Premium Bağlantı Akışı", "13", indent=360),
        toc_line("2.7 Uyum Takibi ve Aktivite Kaydı", "13", indent=360),
        toc_line("2.8 Güvenlik, Yetkilendirme ve Veri İzolasyonu", "13", indent=360),
        toc_line("2.9 Test ve Değerlendirme Yöntemi", "14", indent=360),
        toc_line("2.10 Katmanlar ve Modül Envanteri", "15", indent=360),
        toc_line("ÜÇÜNCÜ BÖLÜM    BULGULAR VE TARTIŞMA", "16", bold=True),
        toc_line("3.1 Prototip Çıktıları", "16", indent=360),
        toc_line("3.2 Normalizasyon Katmanı Bulguları", "19", indent=360),
        toc_line("3.3 Karar Motoru Bulguları", "21", indent=360),
        toc_line("3.4 API ve Erişim Akışı Bulguları", "21", indent=360),
        toc_line("3.5 Literatürle Karşılaştırma", "22", indent=360),
        toc_line("3.6 Tartışma", "23", indent=360),
        toc_line("SONUÇ VE ÖNERİLER", "24", bold=True),
        toc_line("KAYNAKÇA", "25", bold=True),
        toc_line("EKLER", "27", bold=True),
        toc_line("ÖZGEÇMİŞ", "55", bold=True),
    ]
    elems.extend(toc_2)
    elems.append(section_break(footer_rid=roman_footers[5], fmt="decimal", start=1))

    tables = [
        ("G.1 Proje iş-zaman grafiği ve iş paketleri", "3"),
        ("1.1 Literatürdeki çalışmaların karşılaştırılması", "6"),
        ("2.1 Kullanılan yazılım ve teknoloji bileşenleri", "7"),
        ("2.2 Temel veri varlıkları ve görevleri", "9"),
        ("2.3 Eşleştirme motoru karar durumları", "11"),
        ("3.1 Normalizasyon benchmark özet metrikleri", "20"),
        ("3.2 Resolver katmanı dağılımı", "20"),
        ("3.3 Normalizasyon ablation karşılaştırması", "20"),
        ("3.4 Tarif öneri motoru benchmark özet metrikleri", "21"),
        ("3.5 Premium Guard ve Tenant Isolation sonuçları", "21"),
        ("3.6 API endpoint latency sonuçları", "22"),
        ("3.7 Önerilen yaklaşımın literatürle karşılaştırılması", "23"),
        ("3.8 Ara kontrolde görülen test hataları ve final durumu", "22"),
    ]
    table_page = [paragraph("", after=620), paragraph("Tablo Listesi", align="center", bold=True, size=20, after=420)]
    for label, page in tables:
        table_page.append(toc_line(label, page, indent=240))
    elems.extend(table_page)
    elems.append(section_break(footer_rid=roman_footers[6], fmt="decimal", start=1))

    figures = [
        ("2.1 Uçtan uca sistem mimarisi", "8"),
        ("2.2 Çok aşamalı malzeme normalizasyonu", "10"),
        ("2.3 Kural tabanlı karşılaştırma motoru", "12"),
        ("3.1 Mobil uygulamada ana ekran, mutfak ve tarif öneri akışları", "17"),
        ("3.2 Mobil uygulamada plan, dolap, alternatif öğün ve tabak tarama ekranları", "18"),
        ("3.3 Diyetisyen web panelinde dashboard, danışan, plan ve tarif yönetimi ekranları", "18"),
        ("3.4 Web panelde erişim anahtarı, Care Hub, tarif kural girişi ve ayarlar ekranları", "19"),
    ]
    fig_page = [paragraph("", after=620), paragraph("Şekil Listesi", align="center", bold=True, size=20, after=420)]
    for label, page in figures:
        fig_page.append(toc_line(label, page, indent=240))
    elems.extend(fig_page)
    elems.append(section_break(footer_rid=roman_footers[7], fmt="decimal", start=1))

    appendices = [
        ("Ek A", "Demo kontrol listesi"),
        ("Ek B", "Önerilen test verisi örnekleri"),
        ("Ek C", "Final benchmark ve test artefakt dosyaları"),
        ("Ek D", "Backend API uç noktaları"),
        ("Ek E", "Mobil uygulama ekranları"),
        ("Ek F", "Web panel dashboard sayfaları"),
        ("Ek G", "Veritabanı şeması"),
        ("Ek H", "Domain entity sınıfları"),
        ("Ek I", "Test envanteri"),
        ("Ek J", "Mobil API modülleri"),
        ("Ek K", "Web panel API modülleri"),
        ("Ek L", "API endpoint envanteri"),
        ("Ek M", "Backend servis kayıtları"),
        ("Ek N", "docker-compose.yml altyapı özeti"),
        ("Ek O", "Tez odaklı proje özeti ve savunma haritası"),
    ]
    ek_page = [paragraph("", after=620), paragraph("EKLER LİSTESİ", align="center", bold=True, size=22, after=520)]
    for key, label in appendices:
        ek_page.append(paragraph(f"{key}\t{label}", size=20, after=60, tabs=False, left=260))
    elems.extend(ek_page)
    elems.append(section_break(footer_rid=roman_footers[8], fmt="decimal", start=1))
    return elems


def static_footer_xml(text: str) -> bytes:
    xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:ftr xmlns:w="{W}" xmlns:r="{R}">
  <w:p>
    <w:pPr><w:jc w:val="center"/></w:pPr>
    <w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="20"/></w:rPr><w:t>{escape(text)}</w:t></w:r>
  </w:p>
</w:ftr>'''
    return xml.encode("utf-8")


def page_footer_xml(cached: str = "1") -> bytes:
    xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:ftr xmlns:w="{W}" xmlns:r="{R}">
  <w:p>
    <w:pPr><w:jc w:val="center"/></w:pPr>
    <w:r><w:fldChar w:fldCharType="begin"/></w:r>
    <w:r><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>
    <w:r><w:fldChar w:fldCharType="separate"/></w:r>
    <w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="20"/></w:rPr><w:t>{escape(cached)}</w:t></w:r>
    <w:r><w:fldChar w:fldCharType="end"/></w:r>
  </w:p>
</w:ftr>'''
    return xml.encode("utf-8")


def next_footer_names(names: set[str], count: int) -> list[str]:
    n = 1
    result: list[str] = []
    while len(result) < count:
        candidate = f"word/footer{n}.xml"
        if candidate not in names:
            result.append(candidate)
        n += 1
    return result


def next_rids(rels_root: etree._Element, count: int) -> list[str]:
    used = {rel.get("Id") for rel in rels_root.findall("rel:Relationship", namespaces=NS)}
    max_num = 0
    for rid in used:
        m = re.match(r"rId(\d+)$", rid or "")
        if m:
            max_num = max(max_num, int(m.group(1)))
    out: list[str] = []
    n = max_num + 1
    while len(out) < count:
        rid = f"rId{n}"
        if rid not in used:
            out.append(rid)
        n += 1
    return out


def patch_package(path: Path, meta: dict[str, object]) -> None:
    with zipfile.ZipFile(path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    doc_tree = etree.fromstring(files["word/document.xml"])
    rels_tree = etree.fromstring(files["word/_rels/document.xml.rels"])
    content_tree = etree.fromstring(files["[Content_Types].xml"])

    footer_names = next_footer_names(set(files), 10)
    footer_roman_names = footer_names[:9]
    footer_arabic_name = footer_names[9]
    footer_rids = next_rids(rels_tree, 10)
    roman_rids = footer_rids[:9]
    arabic_rid = footer_rids[9]
    rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer"
    for rid, fname in zip(roman_rids, footer_roman_names):
        rels_tree.append(etree.Element(qn("rel:Relationship"), Id=rid, Type=rel_type, Target=fname.replace("word/", "")))
    rels_tree.append(etree.Element(qn("rel:Relationship"), Id=arabic_rid, Type=rel_type, Target=footer_arabic_name.replace("word/", "")))

    existing_overrides = {o.get("PartName") for o in content_tree.findall("ct:Override", namespaces=NS)}
    footer_ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml"
    for fname in (*footer_roman_names, footer_arabic_name):
        part = "/" + fname
        if part not in existing_overrides:
            content_tree.append(etree.Element(qn("ct:Override"), PartName=part, ContentType=footer_ct))

    body = doc_tree.find("w:body", namespaces=NS)
    if body is None:
        raise RuntimeError("document body not found")

    children = list(body)
    start_idx = None
    for i, child in enumerate(children):
        if child.tag == qn("w:p") and norm(text_of_element(child)).upper() == "GİRİŞ":
            start_idx = i
            break
    if start_idx is None:
        raise RuntimeError("GİRİŞ heading not found")

    final_sect = body.find("w:sectPr", namespaces=NS)
    if final_sect is None:
        final_sect = el("w:sectPr")
        body.append(final_sect)

    for old in list(final_sect.findall("w:footerReference", namespaces=NS)):
        final_sect.remove(old)
    for old in list(final_sect.findall("w:pgNumType", namespaces=NS)):
        final_sect.remove(old)
    final_sect.insert(0, el("w:footerReference", {"w:type": "default", "r:id": arabic_rid}))
    pg_num = el("w:pgNumType", {"w:fmt": "decimal", "w:start": "1"})
    pg_mar = final_sect.find("w:pgMar", namespaces=NS)
    if pg_mar is not None:
        final_sect.insert(final_sect.index(pg_mar) + 1, pg_num)
    else:
        final_sect.append(pg_num)

    for child in children[:start_idx]:
        body.remove(child)

    first_body = list(body)[0]
    front = front_matter(meta, roman_rids)
    for element in front:
        body.insert(body.index(first_body), element)

    files["word/document.xml"] = etree.tostring(doc_tree, xml_declaration=True, encoding="UTF-8", standalone=True)
    files["word/_rels/document.xml.rels"] = etree.tostring(rels_tree, xml_declaration=True, encoding="UTF-8", standalone=True)
    files["[Content_Types].xml"] = etree.tostring(content_tree, xml_declaration=True, encoding="UTF-8", standalone=True)
    for fname, label in zip(footer_roman_names, ["vi", "vii", "viii", "ix", "x", "xi", "xii", "xiii", "xiv"]):
        files[fname] = static_footer_xml(label)
    files[footer_arabic_name] = page_footer_xml("1")

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)


def main() -> None:
    source_doc = Document(SRC)
    meta = extract_metadata(source_doc)
    shutil.copyfile(SRC, WORK)
    polish_existing_body(WORK)
    patch_package(WORK, meta)
    shutil.copyfile(WORK, OUT)
    print(OUT)


if __name__ == "__main__":
    main()
