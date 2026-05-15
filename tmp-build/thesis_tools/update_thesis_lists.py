from __future__ import annotations

import json
from pathlib import Path
import re
import unicodedata

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


DOCX = Path(
    r"C:\Users\hy971\source\repos\MyDietitianMobileApp\tmp-build"
    r"\Dytopia_Bitirme_Tezi_Final_Kilavuza_Gore_Duzenlenmis.docx"
)
PAGES_JSON = Path(
    r"C:\Users\hy971\source\repos\MyDietitianMobileApp\tmp-build"
    r"\thesis_tools\render_text.json"
)

FONT = "Times New Roman"


def normalize(text: str) -> str:
    text = unicodedata.normalize("NFKD", text.casefold())
    return "".join(ch for ch in text if ch.isalnum())


def set_rfonts(run) -> None:
    run.font.name = FONT
    rpr = run._element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run._element.insert(0, rpr)
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), FONT)


def clear_paragraph(paragraph) -> None:
    paragraph._p.clear_content()


def write_tabbed(paragraph, title: str, page: str | int, *, indent_cm: float = 0.0) -> None:
    clear_paragraph(paragraph)
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(indent_cm)
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(1)
    fmt.line_spacing = 1
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    try:
        fmt.tab_stops.clear_all()
    except Exception:
        pass
    fmt.tab_stops.add_tab_stop(Cm(15.0 - indent_cm), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = paragraph.add_run(f"{title}\t{page}")
    set_rfonts(run)
    run.font.size = Pt(10.5)


def blank(paragraph) -> None:
    clear_paragraph(paragraph)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(" ")
    set_rfonts(run)
    run.font.size = Pt(1)


def find_page_index_for_title(pages: list[dict], title: str, main_start: int) -> int | None:
    wanted = normalize(title)
    if not wanted:
        return None
    # Prefer exact rendered-line matches so short headings such as "EKLER" do not
    # accidentally match words like "malzemelerle" in body text.
    for page in pages[main_start:]:
        lines = [normalize(line.strip()) for line in page.get("text", "").splitlines() if line.strip()]
        if wanted in lines:
            return page["pageIndex"]
    for page in pages[main_start:]:
        if wanted in normalize(page.get("text", "")):
            return page["pageIndex"]
    # Captions and wrapped lines sometimes lose punctuation; try removing leading labels.
    soft = re.sub(r"^(tablo|sekil|denklem)\s+", "", title, flags=re.I)
    wanted = normalize(soft)
    for page in pages[main_start:]:
        if wanted and wanted in normalize(page.get("text", "")):
            return page["pageIndex"]
    return None


def page_no(pages: list[dict], title: str, main_start: int) -> str:
    idx = find_page_index_for_title(pages, title, main_start)
    if idx is None:
        return ""
    return str(idx - main_start + 1)


def paragraph_range(doc: Document, start_text: str, end_text: str) -> tuple[list, object]:
    texts = [p.text.strip() for p in doc.paragraphs]
    start = texts.index(start_text)
    end = texts.index(end_text)
    return doc.paragraphs[start + 1:end], doc.paragraphs[end]


def insert_before(paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def replace_range(paragraphs: list, end_paragraph, entries: list[tuple[str, str | int, float]]) -> None:
    if len(entries) > len(paragraphs):
        for _ in range(len(entries) - len(paragraphs)):
            paragraphs.append(insert_before(end_paragraph))
    for i, paragraph in enumerate(paragraphs):
        if i < len(entries):
            title, page, indent = entries[i]
            write_tabbed(paragraph, title, page, indent_cm=indent)
        else:
            blank(paragraph)


def main() -> None:
    with PAGES_JSON.open("r", encoding="utf-8") as f:
        payload = json.load(f)
    pages = payload["pages"]

    # Use the physical page where the body heading GİRİŞ is the first visible line.
    main_start = None
    for page in pages:
        lines = [line.strip() for line in page.get("text", "").splitlines() if line.strip()]
        if "GİRİŞ" in lines[:3] and page["pageIndex"] > 5:
            main_start = page["pageIndex"]
            break
    if main_start is None:
        raise RuntimeError("Could not locate GİRİŞ in rendered page text.")

    doc = Document(DOCX)
    texts = [p.text.strip() for p in doc.paragraphs]
    intro_idx = texts.index("GİRİŞ")
    appendix_idx = texts.index("EKLER")
    cv_idx = texts.index("ÖZGEÇMİŞ")

    front_entries = [
        ("ÖZET", "vi", 0.0),
        ("ABSTRACT", "vii", 0.0),
        ("ÖNSÖZ", "viii", 0.0),
        ("KISALTMALAR", "ix", 0.0),
        ("İÇİNDEKİLER", "x", 0.0),
        ("TABLOLAR LİSTESİ", "xii", 0.0),
        ("ŞEKİLLER LİSTESİ", "xiii", 0.0),
        ("EKLER LİSTESİ", "xiv", 0.0),
    ]

    heading_entries: list[tuple[str, str, float]] = []
    for paragraph in doc.paragraphs[intro_idx:]:
        text = paragraph.text.strip()
        style = paragraph.style.name
        if style == "Heading 1":
            if text == "ÖZGEÇMİŞ":
                continue
            heading_entries.append((text, page_no(pages, text, main_start), 0.0))
        elif style == "Heading 2" and not text.startswith("Ek "):
            heading_entries.append((text, page_no(pages, text, main_start), 0.5))

    toc_entries = front_entries + heading_entries
    rng, end = paragraph_range(doc, "İçindekiler", "Tablo Listesi")
    replace_range(rng, end, toc_entries)

    table_entries: list[tuple[str, str, float]] = []
    figure_entries: list[tuple[str, str, float]] = []
    for paragraph in doc.paragraphs[intro_idx:]:
        text = paragraph.text.strip()
        if text.startswith("Tablo "):
            label = re.sub(r"^Tablo\s+", "", text)
            table_entries.append((label, page_no(pages, text, main_start), 0.0))
        elif text.startswith("Şekil "):
            label = re.sub(r"^Şekil\s+", "", text)
            figure_entries.append((label, page_no(pages, text, main_start), 0.0))

    rng, end = paragraph_range(doc, "Tablo Listesi", "Şekil Listesi")
    replace_range(rng, end, table_entries)
    rng, end = paragraph_range(doc, "Şekil Listesi", "EKLER LİSTESİ")
    replace_range(rng, end, figure_entries)

    appendix_entries: list[tuple[str, str, float]] = []
    for paragraph in doc.paragraphs[appendix_idx + 1:cv_idx]:
        text = paragraph.text.strip()
        if paragraph.style.name == "Heading 2" and text.startswith("Ek "):
            label = text.replace(". ", " ", 1)
            appendix_entries.append((label, page_no(pages, text, main_start), 0.0))
    rng, end = paragraph_range(doc, "EKLER LİSTESİ", "GİRİŞ")
    replace_range(rng, end, appendix_entries)

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
