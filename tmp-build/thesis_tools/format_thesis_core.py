from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SRC = Path(
    r"C:\Users\hy971\source\repos\MyDietitianMobileApp\tmp-build"
    r"\Dytopia_Bitirme_Tezi_Final_Referans_Formatli.docx"
)
OUT = Path(
    r"C:\Users\hy971\source\repos\MyDietitianMobileApp\tmp-build"
    r"\Dytopia_Bitirme_Tezi_Final_Kilavuza_Gore_Duzenlenmis.docx"
)

FONT = "Times New Roman"


def set_rfonts(element, name: str = FONT) -> None:
    rpr = element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        element.insert(0, rpr)
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), name)


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = FONT
    set_rfonts(run._element)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_style_font(style, size: float | None, bold: bool | None = None) -> None:
    style.font.name = FONT
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    set_rfonts(style._element)


def clear_paragraph(paragraph) -> None:
    paragraph._p.clear_content()


def add_page_field(paragraph, cached: str = "1") -> None:
    fld_begin = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    fld_begin.append(begin)

    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run.append(instr)

    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)

    text_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), FONT)
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rpr.append(sz)
    text_run.append(rpr)
    text = OxmlElement("w:t")
    text.text = cached
    text_run.append(text)

    fld_end = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    fld_end.append(end)

    paragraph._p.append(fld_begin)
    paragraph._p.append(instr_run)
    paragraph._p.append(separate_run)
    paragraph._p.append(text_run)
    paragraph._p.append(fld_end)


def set_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def set_section_page_numbering(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))
    pg_num.set(qn("w:fmt"), "decimal")


def set_table_cell_margins(table, margin_twips: int = 90) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for side in ("top", "left", "bottom", "right"):
        node = cell_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            cell_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def set_table_width_pct(table, pct: int = 5000) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(pct))


def repeat_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def shade_cell(cell, fill: str = "EDEDED") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def format_paragraph(paragraph, *, size=12, bold=None, align=None, line=1.5,
                     first_indent=True, before=0, after=6, keep_next=False,
                     page_break=False) -> None:
    fmt = paragraph.paragraph_format
    if align is not None:
        paragraph.alignment = align
    fmt.space_before = Pt(before) if before is not None else None
    fmt.space_after = Pt(after) if after is not None else None
    fmt.first_line_indent = Cm(0.7) if first_indent else Cm(0)
    fmt.keep_with_next = keep_next
    fmt.page_break_before = page_break
    if line == 1.0:
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        fmt.line_spacing = 1
    elif line == 1.5:
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        fmt.line_spacing = 1.5
    else:
        fmt.line_spacing = line
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def normalized_sort_key(text: str) -> str:
    repl = str.maketrans({
        "Ç": "C", "ç": "c", "Ğ": "G", "ğ": "g", "İ": "I", "ı": "i",
        "Ö": "O", "ö": "o", "Ş": "S", "ş": "s", "Ü": "U", "ü": "u",
        "Ł": "L", "ł": "l", "Á": "A", "á": "a", "É": "E", "é": "e",
        "Í": "I", "í": "i", "Ó": "O", "ó": "o", "Ú": "U", "ú": "u",
        "’": "'", "“": '"', "”": '"',
    })
    return text.translate(repl).casefold()


def format_document() -> None:
    doc = Document(SRC)

    # Page setup: Tarsus guide target dimensions.
    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(3.0)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)

    styles = doc.styles
    set_style_font(styles["Normal"], 12)
    normal = styles["Normal"].paragraph_format
    normal.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.line_spacing = 1.5
    normal.space_after = Pt(6)
    normal.first_line_indent = Cm(0.7)

    for name, size, align, before, after, page_break in [
        ("Heading 1", 14, WD_ALIGN_PARAGRAPH.CENTER, 0, 12, True),
        ("Heading 2", 12, WD_ALIGN_PARAGRAPH.LEFT, 12, 6, False),
        ("Heading 3", 12, WD_ALIGN_PARAGRAPH.LEFT, 6, 3, False),
    ]:
        style = styles[name]
        set_style_font(style, size, True)
        pf = style.paragraph_format
        pf.alignment = align
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.first_line_indent = Cm(0)
        pf.keep_with_next = True
        pf.page_break_before = page_break

    if "Caption TR" in [s.name for s in styles]:
        caption = styles["Caption TR"]
    else:
        caption = styles.add_style("Caption TR", 1)
    set_style_font(caption, 12, True)
    cpf = caption.paragraph_format
    cpf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cpf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    cpf.line_spacing = 1
    cpf.space_after = Pt(6)
    cpf.first_line_indent = Cm(0)

    texts = [p.text.strip() for p in doc.paragraphs]
    intro_idx = texts.index("GİRİŞ") if "GİRİŞ" in texts else 0
    refs_idx = texts.index("KAYNAKÇA") if "KAYNAKÇA" in texts else len(texts)
    appendix_idx = texts.index("EKLER") if "EKLER" in texts else len(texts)
    cv_idx = texts.index("ÖZGEÇMİŞ") if "ÖZGEÇMİŞ" in texts else len(texts)

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            paragraph.paragraph_format.first_line_indent = Cm(0)
            continue

        if paragraph.style.name.startswith("Heading"):
            if paragraph.style.name == "Heading 1":
                # The guide requires first-degree headings to be uppercase.
                if text != text.upper():
                    paragraph.text = text.upper()
                format_paragraph(
                    paragraph,
                    size=14,
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    line=1.0,
                    first_indent=False,
                    before=0,
                    after=12,
                    keep_next=True,
                    page_break=(idx >= intro_idx),
                )
            else:
                format_paragraph(
                    paragraph,
                    size=12,
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    line=1.5,
                    first_indent=False,
                    before=12 if paragraph.style.name == "Heading 2" else 6,
                    after=6 if paragraph.style.name == "Heading 2" else 3,
                    keep_next=True,
                )
            continue

        if text.startswith(("Tablo ", "Şekil ", "Denklem ")):
            paragraph.style = caption
            format_paragraph(
                paragraph,
                size=12,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                line=1.0,
                first_indent=False,
                before=0,
                after=6,
                keep_next=text.startswith(("Tablo ", "Denklem ")),
            )
            continue

        if idx < intro_idx:
            # Front matter: keep formal centered headings and single-spaced lists.
            if text in {
                "KABUL VE ONAY SAYFASI",
                "ETİK BEYAN",
                "ÖZET",
                "ABSTRACT",
                "ÖNSÖZ",
                "KISALTMALAR",
                "İçindekiler",
                "Tablo Listesi",
                "Şekil Listesi",
                "EKLER LİSTESİ",
            }:
                format_paragraph(
                    paragraph,
                    size=12,
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    line=1.0,
                    first_indent=False,
                    before=0,
                    after=12,
                )
            else:
                align = paragraph.alignment
                if idx < 24:
                    align = WD_ALIGN_PARAGRAPH.CENTER
                elif text.startswith(("API", "B2B", "B2B2C", "CSV", "EF Core", "JWT", "LLM", "MVP", "NRS", "RBAC", "REST", "UI", "UX")):
                    align = WD_ALIGN_PARAGRAPH.LEFT
                format_paragraph(
                    paragraph,
                    size=12,
                    bold=None,
                    align=align,
                    line=1.0,
                    first_indent=False,
                    before=0,
                    after=6,
                )
            continue

        if refs_idx < idx < appendix_idx:
            format_paragraph(
                paragraph,
                size=11,
                bold=None,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                line=1.0,
                first_indent=False,
                before=0,
                after=6,
            )
            paragraph.paragraph_format.left_indent = Cm(1.0)
            paragraph.paragraph_format.first_line_indent = Cm(-1.0)
            continue

        if appendix_idx < idx < cv_idx:
            format_paragraph(
                paragraph,
                size=10,
                bold=None,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                line=1.0,
                first_indent=False,
                before=0,
                after=3,
            )
            continue

        if idx > cv_idx:
            format_paragraph(
                paragraph,
                size=12,
                bold=None,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                line=1.0,
                first_indent=False,
                before=0,
                after=6,
            )
            continue

        if re.match(r"^(Benzerlik\(|Skor\s*=)", text):
            format_paragraph(
                paragraph,
                size=12,
                bold=None,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                line=1.0,
                first_indent=False,
                before=6,
                after=6,
            )
        else:
            format_paragraph(
                paragraph,
                size=12,
                bold=None,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                line=1.5,
                first_indent=True,
                before=0,
                after=6,
            )

    # Figure/image paragraphs.
    for paragraph in doc.paragraphs:
        if paragraph._element.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_after = Pt(3)

    # Keep images inside the new 15 cm text block.
    max_width = Cm(14.8)
    for shape in doc.inline_shapes:
        if shape.width > max_width:
            ratio = shape.height / shape.width
            shape.width = max_width
            shape.height = int(max_width * ratio)

    # Tables: fit to text block, repeat header rows, compact readable type.
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_table_width_pct(table)
        set_table_cell_margins(table)
        if table.rows:
            repeat_header_row(table.rows[0])
            for cell in table.rows[0].cells:
                shade_cell(cell)
        font_size = 9 if len(table.columns) >= 4 else 10
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(table.columns) >= 4 else WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    for run in paragraph.runs:
                        set_run_font(run, size=font_size, bold=(r_idx == 0))

    # Alphabetize bibliography paragraphs and apply hanging indent.
    paragraphs = doc.paragraphs
    texts = [p.text.strip() for p in paragraphs]
    refs_idx = texts.index("KAYNAKÇA") if "KAYNAKÇA" in texts else -1
    appendix_idx = texts.index("EKLER") if "EKLER" in texts else -1
    if refs_idx >= 0 and appendix_idx > refs_idx:
        ref_paras = [p for p in paragraphs[refs_idx + 1:appendix_idx] if p.text.strip()]
        sorted_refs = sorted([p.text.strip() for p in ref_paras], key=normalized_sort_key)
        for p, text in zip(ref_paras, sorted_refs):
            clear_paragraph(p)
            r = p.add_run(text)
            set_run_font(r, size=11)
            format_paragraph(
                p,
                size=11,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                line=1.0,
                first_indent=False,
                before=0,
                after=6,
            )
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-1.0)

    # Main-body numbering: from Giriş onward, Arabic numbers in the upper right.
    main_section = doc.sections[-1]
    main_section.different_first_page_header_footer = False
    main_section.header.is_linked_to_previous = False
    main_section.footer.is_linked_to_previous = False
    main_section.even_page_header.is_linked_to_previous = False
    main_section.even_page_footer.is_linked_to_previous = False
    main_section.first_page_header.is_linked_to_previous = False
    main_section.first_page_footer.is_linked_to_previous = False
    set_section_page_numbering(main_section, start=1)

    for footer in (main_section.footer, main_section.even_page_footer, main_section.first_page_footer):
        for paragraph in footer.paragraphs:
            clear_paragraph(paragraph)
            run = paragraph.add_run(" ")
            set_run_font(run, size=1)
    for header in (main_section.header, main_section.even_page_header, main_section.first_page_header):
        if not header.paragraphs:
            p = header.add_paragraph()
        else:
            p = header.paragraphs[0]
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        add_page_field(p)

    set_update_fields(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    format_document()
