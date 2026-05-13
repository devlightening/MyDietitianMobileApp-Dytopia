from __future__ import annotations

import csv
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
RESULTS = ROOT / "docs" / "thesis-benchmark-results"
INPUT = ROOT / "tmp-build" / "Dytopia_Bitirme_Tezi_Profesyonel_Taslak_v5.docx"
OUTPUT = ROOT / "tmp-build" / "Dytopia_Bitirme_Tezi_Profesyonel_Taslak_v6_benchmark.docx"


def load_json(name: str):
    return json.loads((RESULTS / name).read_text(encoding="utf-8"))


def fmt_num(value, digits: int = 2) -> str:
    return f"{float(value):.{digits}f}".replace(".", ",")


def fmt_ms(value) -> str:
    return f"{fmt_num(value, 4)} ms"


def read_ablation_rows() -> list[list[str]]:
    with (RESULTS / "normalization-ablation.csv").open(encoding="utf-8", newline="") as handle:
        rows = list(csv.DictReader(handle))
    output = []
    for row in rows:
        output.append(
            [
                row["mode"],
                f"%{fmt_num(row['accuracyPct'])}" if row["accuracyPct"] else "Atlandı",
                f"%{fmt_num(row['coveragePct'])}" if row["coveragePct"] else "Atlandı",
                f"%{fmt_num(row['unresolvedRatePct'])}" if row["unresolvedRatePct"] else "Atlandı",
                f"%{fmt_num(row['falseMatchRatePct'])}" if row["falseMatchRatePct"] else "Atlandı",
            ]
        )
    return output


def find_heading(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Heading not found: {prefix}")


def delete_between(start: Paragraph, end: Paragraph) -> None:
    parent = start._element.getparent()
    children = list(parent)
    start_index = children.index(start._element)
    end_index = children.index(end._element)
    for child in children[start_index + 1 : end_index]:
        parent.remove(child)


def paragraph_after(block: Paragraph | Table, text: str = "", style: str | None = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    if isinstance(block, Paragraph):
        block._p.addnext(new_p)
        parent = block._parent
    else:
        block._tbl.addnext(new_p)
        parent = block._parent
    paragraph = Paragraph(new_p, parent)
    if style:
        paragraph.style = style
    paragraph.text = text
    if style == "Normal":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def table_after(block: Paragraph | Table, headers: list[str], rows: list[list[str]]) -> Table:
    parent = block._parent
    doc = parent.part.document
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
            cell.vertical_alignment = 1

    if isinstance(block, Paragraph):
        block._p.addnext(table._tbl)
    else:
        block._tbl.addnext(table._tbl)
    return table


def page_break_after(block: Paragraph | Table) -> Paragraph:
    paragraph = paragraph_after(block, "", "Normal")
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    return paragraph


def caption_after(block: Paragraph | Table, text: str) -> Paragraph:
    try:
        paragraph = paragraph_after(block, text, "Caption TR")
    except Exception:
        paragraph = paragraph_after(block, text, "Normal")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return paragraph


def add_section_32(anchor: Paragraph, norm: dict) -> None:
    current: Paragraph | Table = paragraph_after(
        anchor,
        (
            f"Normalizasyon hattı, canonical eşleşme, alias eşleşmesi, Türkçe karakter/yazım farkı, "
            f"fuzzy eşleşme ve negatif örneklerden oluşan {norm['TotalCases']} senaryoluk test kümesi "
            f"üzerinde gerçek kod çalıştırılarak değerlendirilmiştir. Bu deneyde {norm['CorrectCount']} "
            f"senaryo doğru sonuçlanmış, toplam doğruluk oranı %{fmt_num(norm['AccuracyPct'])} olarak ölçülmüştür. "
            f"Çözülemeyen giriş sayısı {norm['UnresolvedCount']}, unresolved oranı %{fmt_num(norm['UnresolvedRatePct'])}, "
            f"false match oranı ise %{fmt_num(norm['FalseMatchRatePct'])} olarak gerçekleşmiştir."
        ),
    )
    current = caption_after(current, "Tablo 3.1. Normalizasyon benchmark özet metrikleri.")
    current.paragraph_format.page_break_before = True
    current = table_after(
        current,
        ["Metrik", "Değer"],
        [
            ["Toplam senaryo", str(norm["TotalCases"])],
            ["Doğru senaryo", str(norm["CorrectCount"])],
            ["Yanlış/eksik senaryo", str(norm["IncorrectCount"])],
            ["Doğruluk", f"%{fmt_num(norm['AccuracyPct'])}"],
            ["Unresolved oranı", f"%{fmt_num(norm['UnresolvedRatePct'])}"],
            ["False match oranı", f"%{fmt_num(norm['FalseMatchRatePct'])}"],
            ["Ortalama latency", fmt_ms(norm["AverageLatencyMs"])],
            ["Medyan latency", fmt_ms(norm["MedianLatencyMs"])],
            ["P95 latency", fmt_ms(norm["P95LatencyMs"])],
        ],
    )
    current = caption_after(current, "Tablo 3.2. Resolver katmanı dağılımı.")
    current = table_after(
        current,
        ["Katman", "Sayı", "Oran"],
        [
            [layer, str(row["Count"]), f"%{fmt_num(row['Percent'])}"]
            for layer, row in sorted(norm["ResolverLayerDistribution"].items())
        ],
    )
    current = paragraph_after(
        current,
        (
            "Ablation deneyi, katmanların katkısını ayrı ayrı göstermektedir. Canonical-only yaklaşımda "
            "%49,32 olan doğruluk, alias katmanı eklendiğinde %80,82 düzeyine yükselmiş; fuzzy katmanıyla "
            "%98,63 seviyesine ulaşmıştır. OpenAI API anahtarı yapılandırılmadığı için LLM fallback ölçümü "
            "çalıştırılmamış ve bu kısım sonuçlarda açık şekilde skip edilmiştir."
        ),
    )
    current = caption_after(current, "Tablo 3.3. Normalizasyon ablation karşılaştırması.")
    current = table_after(
        current,
        ["Mod", "Doğruluk", "Kapsama", "Unresolved", "False match"],
        read_ablation_rows(),
    )
    paragraph_after(
        current,
        (
            "Pozitif beklenen senaryolarda yalnızca “pirinc pilav” girdisi “Pirinç” malzemesine bağlanamamış "
            "ve unresolved olarak kalmıştır. Bu durum, sistemin belirsiz eşleşmelerde hatalı karar üretmek "
            "yerine güvenli biçimde çözümsüz bırakma davranışını göstermektedir."
        ),
    )


def add_section_33(anchor: Paragraph, recipe: dict) -> None:
    current: Paragraph | Table = paragraph_after(
        anchor,
        (
            f"Tarif öneri motoru; zorunlu, opsiyonel, yasaklı ve alternatif malzeme senaryolarının yanında "
            f"condiment-only guard, boş zorunlu malzeme kalite kontrolü ve opsiyonel eksik durumlarını kapsayan "
            f"{recipe['TotalScenarios']} senaryoluk test kümesi üzerinde değerlendirilmiştir. Motor, "
            f"{recipe['CorrectDecisionCount']} senaryonun tamamında beklenen karar sınıfını üretmiştir."
        ),
    )
    current = caption_after(current, "Tablo 3.4. Tarif öneri motoru benchmark özet metrikleri.")
    current = table_after(
        current,
        ["Metrik", "Değer"],
        [
            ["Toplam senaryo", str(recipe["TotalScenarios"])],
            ["Doğru karar sayısı", str(recipe["CorrectDecisionCount"])],
            ["Recipe Match Accuracy", f"%{fmt_num(recipe['RecipeMatchAccuracyPct'])}"],
            ["Prohibited Filter Success", f"%{fmt_num(recipe['ProhibitedFilterSuccessPct'])}"],
            ["Substitute Handling Success", f"%{fmt_num(recipe['SubstituteHandlingSuccessPct'])}"],
            ["Condiment-only Guard Success", f"%{fmt_num(recipe['CondimentOnlyGuardSuccessPct'])}"],
            ["Ortalama latency", fmt_ms(recipe["AverageLatencyMs"])],
            ["Medyan latency", fmt_ms(recipe["MedianLatencyMs"])],
            ["P95 latency", fmt_ms(recipe["P95LatencyMs"])],
        ],
    )
    paragraph_after(
        current,
        (
            "Bu bulgu, öneri motorunun yalnızca eşleşen malzeme sayısına göre çalışmadığını; yasaklı malzeme, "
            "alternatif malzeme, eksik zorunlu malzeme ve düşük kaliteli condiment-only eşleşmelerini ayrı "
            "karar sınıflarıyla ayırabildiğini göstermektedir. Dolayısıyla kullanıcıya verilen öneri, yalnızca "
            "bir liste değil, gerekçeli ve izlenebilir bir karar çıktısıdır."
        ),
    )


def add_section_34(anchor: Paragraph, premium: dict, api: dict) -> None:
    api_labels = {
        "GET /api/dev/benchmark/acquisition": "Acq.",
        "GET /api/dev/benchmark/hybrid-recipe": "Hybrid",
        "GET /api/dev/benchmark/normalization": "Norm.",
        "GET /api/dev/benchmark/recommendation": "Rec.",
    }
    current: Paragraph | Table = paragraph_after(
        anchor,
        (
            f"Premium Guard ve Tenant Isolation testi {premium['TotalCases']} senaryo ile yürütülmüştür. "
            "Free kullanıcının private tarife erişememesi, premium kullanıcının yalnızca kendi diyetisyeninin "
            "private tariflerini görebilmesi, başka diyetisyenin verilerine erişimin engellenmesi, public fallback, "
            "expired access key ve revoked premium durumları gerçek policy kodu üzerinden sınanmıştır."
        ),
    )
    current = caption_after(current, "Tablo 3.5. Premium Guard ve Tenant Isolation sonuçları.")
    current = table_after(
        current,
        ["Metrik", "Değer"],
        [
            ["Toplam senaryo", str(premium["TotalCases"])],
            ["Doğru senaryo", str(premium["CorrectCases"])],
            ["Premium Guard Success", f"%{fmt_num(premium['PremiumGuardSuccessPct'])}"],
            ["Tenant Isolation Success", f"%{fmt_num(premium['TenantIsolationSuccessPct'])}"],
            ["Başarısız senaryo", "Yok" if not premium["FailedCases"] else ", ".join(premium["FailedCases"])],
        ],
    )
    current = paragraph_after(
        current,
        (
            "API gecikme ölçümleri ASP.NET Core test-server üzerinde HTTP çağrılarıyla yapılmıştır. Bu ölçüm, "
            "routing, controller çalışması ve JSON serileştirme maliyetlerini içerir; mobil ağ, internet ve "
            "harici servis gecikmelerini içermez. Her endpoint için 30 tekrar yapılmış ve tüm endpointlerde hata "
            "sayısı 0 olarak kaydedilmiştir."
        ),
    )
    current = caption_after(current, "Tablo 3.6. API endpoint latency sonuçları.")
    current = table_after(
        current,
        ["Akış", "n", "Ort. ms", "Med. ms", "P95 ms", "Hata"],
        [
            [
                api_labels.get(operation, operation.replace("GET /api/dev/benchmark/", "")),
                str(row["Count"]),
                fmt_num(row["AverageMs"], 2),
                fmt_num(row["MedianMs"], 2),
                fmt_num(row["P95Ms"], 2),
                str(row["ErrorCount"]),
            ]
            for operation, row in sorted(api.items())
        ],
    )
    paragraph_after(
        current,
        (
            "Mevcut API test paketi 209 testten oluşmuş; 200 test geçmiş, 2 test başarısız olmuş ve 7 test "
            "atlanmıştır. Benchmark artifact testi, benchmark endpoint smoke testleri ve API latency artifact "
            "testi başarıyla tamamlanmıştır. OpenAI API anahtarı yapılandırılmadığı için OpenAI fallback "
            "senaryoları çalıştırılmamış; bu sonuçlar deterministik normalizasyon, tarif öneri ve erişim "
            "kontrol katmanlarının gerçek çalıştırma çıktılarıyla sınırlı tutulmuştur."
        ),
    )


def update_page_count(doc: Document, pages: int) -> None:
    for paragraph in doc.paragraphs:
        if "Haziran 2026," in paragraph.text and ("sayfa" in paragraph.text or "pages" in paragraph.text):
            for run in paragraph.runs:
                run.text = re.sub(r"Haziran 2026, \d+ sayfa", f"Haziran 2026, {pages} sayfa", run.text)
                run.text = re.sub(r"Haziran 2026, \d+ pages", f"Haziran 2026, {pages} pages", run.text)


def update_table_list_and_caption_numbers(doc: Document) -> None:
    table_list_entries = [
        "Tablo 3.1. Normalizasyon benchmark özet metrikleri ........................................",
        "Tablo 3.2. Resolver katmanı dağılımı ........................................",
        "Tablo 3.3. Normalizasyon ablation karşılaştırması ........................................",
        "Tablo 3.4. Tarif öneri motoru benchmark özet metrikleri ........................................",
        "Tablo 3.5. Premium Guard ve Tenant Isolation sonuçları ........................................",
        "Tablo 3.6. API endpoint latency sonuçları ........................................",
        "Tablo 3.7. Önerilen yaklaşımın literatürle karşılaştırılması ........................................",
    ]
    front_table_items = [
        paragraph
        for paragraph in doc.paragraphs[:140]
        if paragraph.text.strip().startswith("Tablo 3.")
    ]
    if len(front_table_items) >= 3:
        for paragraph, text in zip(front_table_items[:3], table_list_entries[:3]):
            paragraph.text = text
        current: Paragraph | Table = front_table_items[2]
        for text in table_list_entries[3:]:
            current = paragraph_after(current, text, "Normal")

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Tablo 3.3. Önerilen yaklaşımın literatürle karşılaştırılması"):
            paragraph.text = "Tablo 3.7. Önerilen yaklaşımın literatürle karşılaştırılması."


def main() -> None:
    doc = Document(INPUT)
    norm = load_json("normalization-summary.json")
    recipe = load_json("recipe-engine-summary.json")
    premium = load_json("premium-guard-summary.json")
    api = load_json("api-latency-summary.json")

    h32 = find_heading(doc, "3.2.")
    h33 = find_heading(doc, "3.3.")
    h34 = find_heading(doc, "3.4.")
    h35 = find_heading(doc, "3.5.")

    delete_between(h34, h35)
    add_section_34(h34, premium, api)

    delete_between(h33, h34)
    add_section_33(h33, recipe)

    delete_between(h32, h33)
    add_section_32(h32, norm)

    update_table_list_and_caption_numbers(doc)
    update_page_count(doc, 82)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
