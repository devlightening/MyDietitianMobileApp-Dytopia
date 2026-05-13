from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:/Users/hy971/source/repos/MyDietitianMobileApp")
DOCX_PATH = ROOT / "tmp-build" / "Dytopia_Bitirme_Tezi_Profesyonel_Taslak_v2.docx"
BACKUP_PATH = ROOT / "tmp-build" / "Dytopia_Bitirme_Tezi_Profesyonel_Taslak_v2_backup_once_v4.docx"
OUT_PATH = ROOT / "tmp-build" / "Dytopia_Bitirme_Tezi_Profesyonel_Taslak_v4.docx"


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_paragraph(doc: Document, text: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    return None


def find_first_startswith(doc: Document, prefix: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    return None


def has_heading(doc: Document, heading: str) -> bool:
    return any(paragraph.text.strip() == heading for paragraph in doc.paragraphs)


def set_paragraph_text(paragraph: Paragraph | None, text: str) -> None:
    if paragraph is None:
        return
    style = paragraph.style
    paragraph.text = text
    paragraph.style = style


def add_block(anchor: Paragraph, heading: str, paragraphs: list[str], heading_style: str = "Heading 2") -> Paragraph:
    current = insert_paragraph_after(anchor, heading, heading_style)
    for text in paragraphs:
        current = insert_paragraph_after(current, text, "Normal")
    return current


def add_bullets(anchor: Paragraph, items: list[str]) -> Paragraph:
    current = anchor
    for item in items:
        current = insert_paragraph_after(current, f"• {item}", "Normal")
    return current


def count_files(path: Path, pattern: str) -> int:
    if not path.exists():
        return 0
    return len(list(path.rglob(pattern)))


def count_direct_files(path: Path, pattern: str) -> int:
    if not path.exists():
        return 0
    return len(list(path.glob(pattern)))


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(str(DOCX_PATH))

    for section in doc.sections:
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    set_paragraph_text(
        find_first_startswith(doc, "Görüntüden malzeme tanıma ve LLM tabanlı"),
        (
            "OpenAI destekli metin ve görsel çözümleme, proje kapsamında yardımcı ve kontrollü bir katman "
            "olarak ele alınmıştır. Tezin ana bilimsel katkısı, yapay zekâ çıktısını tek başına karar kabul etmek "
            "değil; malzeme normalizasyonu, taksonomi ve kural tabanlı öneri ile diyetisyen-danışan "
            "senkronizasyonunu denetlenebilir bir karar destek hattında birleştirmektir. Bu nedenle görsel "
            "tarama, market fişi okuma, tabak fotoğrafı analizi ve günlük oyun içeriği üretimi sistemin "
            "bütünsel mimarisinde anlatılmış; nicel değerlendirme odağı ise resolver başarısı, öneri motoru "
            "doğruluğu, erişim güvenliği ve API davranışı üzerinde tutulmuştur."
        ),
    )

    set_paragraph_text(
        find_first_startswith(doc, "Bu çalışmanın birincil özgün katkısı;"),
        (
            "Bu çalışmanın birincil özgün katkısı; serbest metin, barkod, fiş veya görsel kaynaklı gürültülü "
            "malzeme girdilerini çok katmanlı bir resolver hattı üzerinden standart IngredientId kimliklerine "
            "dönüştürmesi ve bu standart veriyi açıklanabilir tarif önerisi için kullanmasıdır. Resolver zinciri "
            "canonical exact match, alias match, önceden onaylanmış vision label mapping, provisional mapping, "
            "fuzzy match, kontrollü LLM fallback ve unresolved çıktılarından oluşur. Böylece sistem, yalnızca "
            "metin benzerliğine dayanan basit bir arama mekanizması yerine; taksonomi, kullanıcı onayı, kayıt "
            "altına alma ve kural tabanlı doğrulama katmanlarını birlikte kullanan izlenebilir bir karar destek "
            "yaklaşımı sunar."
        ),
    )

    set_paragraph_text(
        find_first_startswith(doc, "Bu model, her klinik için ayrı mobil uygulama"),
        (
            "Access Key temelli çoklu kiracılı yapı, klinik veya diyetisyen bazlı ayrı istemci uygulamaları "
            "yayınlama ihtiyacını azaltırken danışan verisinin aktif diyetisyen bağlantısı üzerinden "
            "yetkilendirilmesini sağlar. Tek mobil uygulama, doğrulanan erişim anahtarına göre bağlı "
            "diyetisyenin plan, tarif, mesajlaşma ve marka bilgilerini gösterir. Bu kurgu, tez kapsamında "
            "rol tabanlı yetkilendirme, veri izolasyonu ve sürdürülebilir bakım maliyeti açısından "
            "değerlendirilmiştir."
        ),
    )

    set_paragraph_text(
        find_first_startswith(doc, "Geliştirilen prototip, danışan mobil uygulaması"),
        (
            "Geliştirilen prototip üç ana katmandan oluşmaktadır: danışanın günlük kullanımını sağlayan Expo "
            "React Native mobil uygulama, diyetisyenin yönetim ve takip süreçlerini yürüttüğü Next.js web "
            "panel ve iş kurallarını/veri erişimini yöneten .NET 8 backend API. Mobil tarafta welcome, login, "
            "dashboard, today, kitchen, ingredient scan, barcode scan, receipt scan, meal log, pantry, shopping "
            "list, plans, favorites, game center, profile, measurements, notifications ve messages ekranları "
            "bulunmaktadır. Web panel tarafında dashboard, danışan listesi, danışan detayı, plan oluşturma, "
            "tarif oluşturma ve içe aktarma, recipe match, access key, care hub, branding, admin ingredient "
            "yönetimi ve owner panel modülleri yer almaktadır."
        ),
    )

    set_paragraph_text(
        find_first_startswith(doc, "Prototip, tez savunmasında gösterilebilir somut bir ürün"),
        (
            "Backend tarafında JWT kimlik doğrulama, rol temelli yetkilendirme, PostgreSQL/EF Core veri modeli, "
            "SignalR canlı veri akışı, OpenAI ping ve vision ping uç noktaları, tarif öneri motoru, malzeme "
            "normalizasyonu, uyum takibi, oyunlaştırma ve meal log AI analizi birlikte çalışmaktadır. Bu yapı, "
            "tezi yalnızca ekran geliştirme çalışması olmaktan çıkarıp veri standardizasyonu ve açıklanabilir "
            "karar desteği problemi olarak konumlandırır."
        ),
    )

    set_paragraph_text(
        find_first_startswith(doc, "Premium aktivasyon akışı, tez projesinin SaaS"),
        (
            "Premium aktivasyon akışı, sistemin çok kullanıcılı SaaS yapısını somutlaştıran temel bulgulardan "
            "biridir. Danışan Access Key girmeden Free modda kalır ve public tarif havuzunu kullanır. Kod "
            "doğrulandığında kullanıcı Premium moda geçer, activeDietitianId atanır ve bağlı diyetisyenin özel "
            "tarifleri, planları, mesajları ve uyum takip verileri görünür hâle gelir. Bu akış hem istemci "
            "tarafındaki guard kontrolleriyle hem de backend policy ve bağlantı süresi denetimleriyle "
            "korunmuştur."
        ),
    )

    set_paragraph_text(
        find_first_startswith(doc, "Bununla birlikte sistemin bazı sınırlılıkları vardır."),
        (
            "Bununla birlikte sistemin bazı sınırlılıkları vardır. Malzeme sözlüğünün kapsamı, alias kalitesi ve "
            "fuzzy eşik değerleri normalizasyon başarısını doğrudan etkiler. OpenAI destekli metin ve görsel "
            "çözümleme belirsiz girdilerde yararlı olsa da nihai karar verici değildir; görselden çıkan etiketler "
            "yeniden Ingredients tablosuna çözümlenir ve güven skoru, kullanıcı onayı ve loglama ile denetlenir. "
            "Ayrıca tabak fotoğrafından kalori/makro tahmini yaklaşık bilgi üretir; klinik tanı, tedavi veya "
            "kesin beslenme reçetesi yerine diyetisyen kontrolündeki yardımcı veri olarak değerlendirilmelidir."
        ),
    )

    if not has_heading(doc, "2.1.2. Katmanlı Sistem Mimarisi ve Bileşen Rolleri"):
        anchor = find_paragraph(doc, "Şekil 2.1. Uçtan uca sistem mimarisi.")
        if anchor is not None:
            add_block(
                anchor,
                "2.1.2. Katmanlı Sistem Mimarisi ve Bileşen Rolleri",
                [
                    (
                        "Backend katmanı; MyDietitianMobileApp.Api, Application, Domain ve Infrastructure "
                        "projeleriyle ayrıştırılmıştır. Api projesi HTTP uç noktalarını, JWT doğrulamasını, "
                        "Swagger yapılandırmasını, SignalR hub'larını ve OpenAI sağlık kontrollerini barındırır. "
                        "Application katmanı iş akışları, command/query handler'ları, oyunlaştırma ve günlük "
                        "içerik üretim mantığını; Domain katmanı entity, enum ve servis sözleşmelerini; "
                        "Infrastructure katmanı ise EF Core repository'leri, PostgreSQL context'i, OpenAI/Ollama "
                        "servisleri, vision servisleri, import parser'ları ve tarif öneri motorunu içerir."
                    ),
                    (
                        "Web panel, diyetisyenin sistemdeki kural ve içerik üretim arayüzüdür. Diyetisyen bu "
                        "panel üzerinden danışanlarını, diyet planlarını, tarif havuzunu, access key üretimini, "
                        "care hub mesajlaşmasını, branding ayarlarını ve admin ingredient yönetimini yürütür. "
                        "Bu nedenle panel, yalnızca yönetim ekranı değil; diyetisyen bilgisinin yapılandırılmış "
                        "kurallara dönüştürüldüğü üretim katmanı olarak ele alınmıştır."
                    ),
                    (
                        "Mobil uygulama, danışanın veri toplama ve karar desteği deneyimini üstlenir. Danışan "
                        "malzeme ekleyebilir, barkod/fiş/görsel tarayabilir, mutfak önerisi alabilir, öğün kaydı "
                        "oluşturabilir, alışveriş listesini ve pantry bilgisini yönetebilir, diyetisyen planlarını "
                        "görüntüleyebilir ve oyunlaştırma modülleriyle günlük bağlılığını sürdürebilir."
                    ),
                ],
            )

    if not has_heading(doc, "2.3.1. OpenAI Destekli Kontrollü Çözümleme"):
        anchor = find_first_startswith(doc, "Canonical eşleşme, en güvenilir")
        if anchor is not None:
            add_block(
                anchor,
                "2.3.1. OpenAI Destekli Kontrollü Çözümleme",
                [
                    (
                        "OpenAI entegrasyonu sistemde dört ana yardımcı görev için kullanılmıştır: serbest "
                        "metin malzeme normalizasyonunda LLM fallback, GPT-4o Vision ile mutfak veya market "
                        "fişi görselinden malzeme çıkarımı, tabak fotoğrafından yaklaşık öğün analizi ve günlük "
                        "mini oyun içeriklerinin üretimi. Bu görevler OpenAiIngredientLlmClient, "
                        "VisionIngredientService, ClientMealLogController ve DailyGameContentGenerator gibi "
                        "bileşenler üzerinden uygulanmıştır."
                    ),
                    (
                        "OpenAI çıktıları sistemde doğrudan nihai karar olarak kabul edilmez. Metin "
                        "normalizasyonunda model yalnızca sınırlı aday listesi içinden seçim yapar; görselden "
                        "çıkan etiketler yeniden Ingredients tablosuna çözümlenir; meal log analizleri kullanıcıya "
                        "yaklaşık bilgi olarak sunulur; günlük oyun üretiminde ise OpenAI kullanılamadığında "
                        "fallback içerik devreye girer. API anahtarı OpenAI:ApiKey veya OPENAI_API_KEY üzerinden "
                        "çözülür ve loglara yazılmaz."
                    ),
                    (
                        "Bu tasarım, yapay zekâ bileşenini karar verici değil, deterministik katmanların "
                        "çözemediği belirsiz girdilerde aday üreten denetimli yardımcı bileşen olarak "
                        "konumlandırır. JSON yanıt biçimi, aday kısıtlama, güven skoru, hallucination guard, "
                        "kullanıcı onayı ve işlem logları bu denetimi destekleyen başlıca mekanizmalardır."
                    ),
                ],
            )

    if not has_heading(doc, "2.5.1. Tarif Öneri Motorunda Kalite Koruma Kuralları"):
        anchor = find_paragraph(doc, "Denklem 2.1. Tarif eşleştirme için ağırlıklı uygunluk skoru.")
        if anchor is not None:
            add_block(
                anchor,
                "2.5.1. Tarif Öneri Motorunda Kalite Koruma Kuralları",
                [
                    (
                        "Tarif öneri motoru, yalnızca isim benzerliği veya malzeme sayısı üzerinden öneri "
                        "üretmez. RecipeRecommendationEngine içinde zorunlu malzeme varlığı, opsiyonel "
                        "malzeme katkısı, yasaklı malzeme çakışması, alternatif/ikame kullanımı ve tarif "
                        "kalite kontrolleri birlikte değerlendirilir. Yasaklı malzeme tespit edildiğinde tarif "
                        "skorlanmadan elenir."
                    ),
                    (
                        "Sistem, yalnızca tuz, yağ, sos gibi yardımcı bileşenlere dayalı anlamsız eşleşmelerin "
                        "yüksek skor almasını engellemek için condiment-only guard yaklaşımını kullanır. Ayrıca "
                        "zorunlu malzemesi boş olan tariflerin hatalı biçimde yüzde yüz uyumlu görünmemesi için "
                        "ek kalite kontrol uygulanır. Bu sayede öneri sonucu kullanıcıya hem uygulanabilir hem de "
                        "açıklanabilir gerekçelerle sunulur."
                    ),
                    (
                        "Premium kullanıcılar için öneri havuzu bağlı diyetisyenin özel tarifleri ve uygun public "
                        "fallback politikasıyla sınırlandırılır. Böylece tarif önerisi yalnızca teknik eşleşmeye "
                        "değil, diyetisyen-danışan ilişkisine ve yetkilendirme bağlamına da bağlı hâle gelir."
                    ),
                ],
            )

    metrics_anchor = find_paragraph(doc, "Tablo 3.2. Karar motoru örnek test senaryoları.")
    if metrics_anchor is not None and not any("Normalization Accuracy" in p.text for p in doc.paragraphs):
        current = insert_paragraph_after(
            metrics_anchor,
            (
                "Değerlendirme bölümünde özellikle şu metrikler raporlanmalıdır: Normalization Accuracy, "
                "Resolver Layer Distribution, Unresolved Rate, False Match Rate, Recipe Match Accuracy, "
                "Prohibited Filter Success, Average API Latency, OpenAI Fallback Usage ve Premium Guard Success. "
                "Bu metrikler, sistemin yalnızca çalıştığını değil; hangi katmanda ne kadar doğru, izlenebilir "
                "ve güvenli çalıştığını göstermeyi amaçlar."
            ),
            "Normal",
        )
        add_bullets(
            current,
            [
                "Normalization Accuracy: Kullanıcı girdisinin doğru IngredientId değerine bağlanma oranı.",
                "Resolver Layer Distribution: Canonical, alias, mapping, fuzzy ve LLM katmanlarının kullanım dağılımı.",
                "Unresolved Rate: Sistem tarafından güvenli biçimde çözülemeyen girdilerin oranı.",
                "Recipe Match Accuracy: Beklenen tarif karar sınıfının doğru üretilme oranı.",
                "Prohibited Filter Success: Yasaklı malzeme içeren tariflerin doğru elenme oranı.",
                "Average API Latency: Kitchen match akışının ortalama ve p95 yanıt süresi.",
                "OpenAI Fallback Usage: OpenAI'nin yalnızca gerekli belirsiz senaryolarda devreye girme oranı.",
                "Premium Guard Success: Free kullanıcının private tarif veya plana erişememe doğruluğu.",
            ],
        )

    if not has_heading(doc, "Ek O. Tez Odaklı Proje Özeti ve Savunma Haritası"):
        anchor = find_first_startswith(doc, "• Danışman tarafından istenen özel format")
        if anchor is not None:
            current = add_block(
                anchor,
                "Ek O. Tez Odaklı Proje Özeti ve Savunma Haritası",
                [
                    (
                        "Dytopia / MyDietitianMobileApp; diyetisyen destekli mobil beslenme platformu olarak "
                        "tasarlanmıştır. Sistem .NET 8 backend, Next.js web panel, Expo React Native mobil "
                        "uygulama ve PostgreSQL veri tabanından oluşur. Tezin teknik odağı, kullanıcıların "
                        "serbest metin, barkod, fiş veya görsel ile ürettiği dağınık beslenme verisini standart "
                        "malzeme kimliklerine dönüştürmek ve bu standart verilerle kural tabanlı tarif önerisi "
                        "üretmektir."
                    ),
                    (
                        "Savunmada projenin yalnızca mobil uygulama geliştirme çalışması olmadığı özellikle "
                        "vurgulanmalıdır. Ana mühendislik problemi; hatalı, eksik, Türkçe karakter içermeyen veya "
                        "belirsiz malzeme girdilerinin güvenilir biçimde çözümlenmesi; bu çözümlenen verinin "
                        "zorunlu, opsiyonel, yasak ve alternatif tarif kurallarıyla karşılaştırılması; sonuçların "
                        "diyetisyen ve danışan arasında izlenebilir hâle getirilmesidir."
                    ),
                    (
                        "OpenAI kullanımı sistemde kontrollü yardımcı katmandır. OpenAI metin normalizasyonu, "
                        "vision tabanlı malzeme/fiş okuma, tabak fotoğrafı analizi ve günlük oyun içeriği üretimi "
                        "için kullanılabilir; ancak nihai öneri kararı deterministic resolver, Ingredients tablosu, "
                        "güven skoru, kullanıcı onayı, loglama ve kural tabanlı tarif motoru tarafından denetlenir."
                    ),
                ],
            )
            current = insert_paragraph_after(current, "Savunmada öne çıkarılacak ana iddialar:", "Normal")
            add_bullets(
                current,
                [
                    "Kullanıcı malzemeleri çok katmanlı resolver zinciriyle standart IngredientId kimliklerine dönüştürülmektedir.",
                    "Tarif önerileri zorunlu, opsiyonel, yasak ve alternatif malzeme rollerine göre açıklanabilir biçimde üretilmektedir.",
                    "Diyetisyen web paneli, danışan mobil uygulaması ve backend API aynı veri modeli üzerinde senkronize çalışmaktadır.",
                    "Access Key yapısı, premium bağlantı ve tenant isolation mantığını uygulama düzeyinde göstermektedir.",
                    "OpenAI nihai karar verici değil; belirsiz metin/görsel girdilerde kontrollü yardımcı bileşendir.",
                    "Değerlendirme, normalizasyon doğruluğu, öneri motoru başarısı, erişim koruması ve API gecikmesi metrikleriyle yapılmalıdır.",
                ],
            )

    controller_count = count_direct_files(ROOT / "src/MyDietitianMobileApp.Api/Controllers", "*.cs")
    mobile_screen_count = count_direct_files(ROOT / "mobile-app/src/screens", "*.tsx")
    web_page_count = count_files(ROOT / "web-panel/app/dashboard", "page.tsx")
    test_file_count = count_files(ROOT / "tests/MyDietitianMobileApp.Api.Tests", "*.cs")

    summary_text = (
        f"Proje taramasında backend tarafında {controller_count} controller dosyası, mobil uygulamada "
        f"{mobile_screen_count} ekran dosyası, web panel dashboard altında {web_page_count} sayfa ve "
        f"Api.Tests projesinde {test_file_count} test/yardımcı C# dosyası belirlenmiştir. Bu envanter, "
        "ekler bölümünde modül bazlı olarak sunulmuş ve tezin kapsamının kaynak kodla uyumlu olduğunu "
        "göstermek için kullanılmıştır."
    )
    if not any("Proje taramasında backend tarafında" in p.text for p in doc.paragraphs):
        anchor = find_first_startswith(doc, "Backend tarafında JWT kimlik")
        if anchor is not None:
            insert_paragraph_after(anchor, summary_text, "Normal")

    doc.save(str(DOCX_PATH))
    doc.save(str(OUT_PATH))
    print(DOCX_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
