from __future__ import annotations

import os
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
SOURCE_DOCX = ROOT / "tmp-build" / "son_Dytopia_Bitirme_Tezi_Profesyonel_Taslak_v6_benchmark.docx"
OUT_DOCX = ROOT / "tmp-build" / "Dytopia_Bitirme_Tezi_Final_v7.docx"
REPORT_DIR = ROOT / "docs" / "thesis-finalization-v6-body"
ASSET_DIR = REPORT_DIR / "assets"
PAGE_COUNT = os.environ.get("THESIS_FINAL_PAGE_COUNT")


def paragraph_text(paragraph: Paragraph) -> str:
    return " ".join((paragraph.text or "").split())


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def insert_paragraph_before(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def cm_to_twips(width_cm: float) -> int:
    return int(width_cm * 567)


def set_cell_margins(cell, top: int = 80, start: int = 80, bottom: int = 80, end: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(cm_to_twips(width_cm)))
    tc_w.set(qn("w:type"), "dxa")


def mark_row_no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def set_table_layout(table: Table, widths_cm: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(cm_to_twips(sum(widths_cm))))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        mark_row_no_split(row)
        for idx, width in enumerate(widths_cm):
            set_cell_width(row.cells[idx], width)


def format_table(
    table: Table,
    widths_cm: list[float],
    body_font_pt: float = 9.5,
    header_font_pt: float = 9.5,
    left_columns: set[int] | None = None,
) -> None:
    left_columns = left_columns or set()
    set_table_layout(table, widths_cm)
    mark_header_row(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=95, start=95, bottom=95, end=95)
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if col_index in left_columns and row_index > 0 else WD_ALIGN_PARAGRAPH.CENTER
                )
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.size = Pt(header_font_pt if row_index == 0 else body_font_pt)
                    if row_index == 0:
                        run.bold = True


def insert_table_after(paragraph: Paragraph, rows: list[list[str]], widths_cm: list[float] | None = None) -> Table:
    doc = paragraph.part.document
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for idx, value in enumerate(rows[0]):
        table.rows[0].cells[idx].text = value
        for run in table.rows[0].cells[idx].paragraphs[0].runs:
            run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    if widths_cm:
        format_table(table, widths_cm)
    paragraph._p.addnext(table._tbl)
    return table


def insert_paragraph_after_table(table: Table, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    new_paragraph = Paragraph(new_p, table._parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def find_paragraph(doc: Document, exact: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph_text(paragraph) == exact:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact}")


def find_startswith(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph_text(paragraph).startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph prefix not found: {prefix}")


def find_last_startswith(doc: Document, prefix: str) -> Paragraph:
    for paragraph in reversed(doc.paragraphs):
        if paragraph_text(paragraph).startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph prefix not found: {prefix}")


def find_contains(doc: Document, *tokens: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        text = paragraph_text(paragraph)
        if all(token in text for token in tokens):
            return paragraph
    raise ValueError(f"Paragraph containing tokens not found: {tokens}")


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def delete_body_range(doc: Document, start_prefix: str, end_prefix: str) -> None:
    start_paragraph = find_startswith(doc, start_prefix)
    end_paragraph = find_startswith(doc, end_prefix)
    body = doc._body._element
    children = list(body)
    start_index = children.index(start_paragraph._element)
    end_index = children.index(end_paragraph._element)
    if end_index <= start_index:
        return
    for child in children[start_index:end_index]:
        body.remove(child)


def collect_between(doc: Document, start_exact: str, end_exact: str) -> list[Paragraph]:
    paragraphs = list(doc.paragraphs)
    start = next(i for i, p in enumerate(paragraphs) if paragraph_text(p) == start_exact)
    end = next(i for i, p in enumerate(paragraphs[start + 1 :], start=start + 1) if paragraph_text(p) == end_exact)
    return paragraphs[start + 1 : end]


def update_front_matter(doc: Document) -> None:
    if not PAGE_COUNT:
        return
    for paragraph in doc.paragraphs:
        text = paragraph_text(paragraph)
        if text.startswith("Lisans Bitirme Tezi, Bilgisayar Mühendisliği Bölümü"):
            set_paragraph_text(
                paragraph,
                "Lisans Bitirme Tezi, Bilgisayar Mühendisliği Bölümü "
                f"Danışman: Dr. Öğr. Üyesi Volkan ATEŞ Haziran 2026, {PAGE_COUNT} sayfa",
            )
        elif text.startswith("Graduation Thesis, Department of Computer Engineering"):
            set_paragraph_text(
                paragraph,
                "Graduation Thesis, Department of Computer Engineering "
                f"Supervisor: Dr. Öğr. Üyesi Volkan ATEŞ June 2026, {PAGE_COUNT} pages",
            )


def update_core_text(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph_text(paragraph)
        if text.startswith("Normalizasyon için altın standart bir veri seti tanımlanır."):
            set_paragraph_text(
                paragraph,
                "Bu çalışma kapsamında normalizasyon hattı; canonical eşleşme, alias eşleşmesi, "
                "Türkçe karakter/yazım farkı, fuzzy eşleşme ve negatif örnekleri içeren kontrollü "
                "73 senaryoluk bir altın standart test kümesi üzerinde değerlendirilmiştir. Test "
                "kümesi, gerçek kullanıcı girdilerinde beklenen temel hata türlerini temsil edecek "
                "şekilde hazırlanmıştır. Daha geniş ve gerçek kullanıcı kaynaklı veri setleriyle "
                "değerlendirme gelecek çalışma kapsamında ele alınmıştır.",
            )
        elif text.startswith("Veri setinin temsil gücü, alias sözlüğünün kapsamı"):
            set_paragraph_text(
                paragraph,
                "Normalizasyon deneyinde kullanılan 73 senaryoluk test kümesi; temel yazım "
                "hataları, alias ifadeleri, Türkçe karakter farklılıkları ve negatif örnekleri "
                "kapsasa da gerçek kullanıcı davranışlarının tamamını temsil etmez. Bu nedenle "
                "sonuçlar kontrollü senaryo tabanlı doğrulama olarak değerlendirilmelidir. "
                "Alias sözlüğünün kapsamı ve fuzzy eşik değerleri sonuçları doğrudan etkileyebilir. "
                "LLM fallback kullanımı ise nihai karar verici olarak değil, belirsiz girdilerde "
                "kontrollü yardımcı katman olarak ele alınmıştır.",
            )
        elif text.startswith("Mevcut API test paketi 209 testten oluşmuş;"):
            set_paragraph_text(
                paragraph,
                "Final doğrulamada Api.Tests paketi tekrar çalıştırılmıştır. İlk ara kontrolde "
                "başarısız görünen barkod çözümleme ve multimodal benchmark senaryoları analiz "
                "edilmiş; offline barkod fallback davranışı ile vision closed-set benchmark "
                "konfigürasyonu düzeltilmiştir. Final koşuda 210 testten 203'ü geçmiş, 0 test "
                "başarısız olmuş ve 7 test atlanmıştır. OpenAI API anahtarı yapılandırılmadığı "
                "için LLM fallback senaryoları çalıştırılmamış; raporlanan ana metrikler "
                "deterministik normalizasyon, tarif öneri ve erişim kontrol katmanlarının gerçek "
                "çalıştırma çıktılarıyla sınırlı tutulmuştur.",
            )
        elif text.startswith("Çalışmanın akademik çıktısı,"):
            set_paragraph_text(
                paragraph,
                "Çalışmanın akademik çıktısı, “ingredient normalization + taxonomy-aware "
                "rule-based recommendation + compliance tracking” üçlüsünün diyetisyen destekli "
                "mobil sağlık bağlamında uygulanmasıdır. Deneysel bulgular, geliştirilen sistemin "
                "temel karar destek bileşenlerinin beklenen şekilde çalıştığını göstermiştir. "
                "Normalizasyon katmanı 73 temsilî senaryoda %98,63 doğruluk ve %0,00 false match "
                "oranı sağlamış; tarif öneri motoru 36 karar senaryosunun tamamında doğru "
                "sınıflandırma üretmiştir. Premium Guard ve Tenant Isolation testlerinde 10 "
                "senaryonun tamamı başarılı sonuçlanmıştır. Bu sonuçlar, Dytopia prototipinin "
                "yalnızca arayüz tabanlı bir mobil uygulama değil, ölçülebilir ve açıklanabilir "
                "bir karar destek sistemi olduğunu desteklemektedir.",
            )
        elif text.strip().startswith("POSTGRES_PASSWORD:"):
            set_paragraph_text(paragraph, "      POSTGRES_PASSWORD: [demo ortamı için maskelendi]")
        elif text.strip().startswith("PGADMIN_DEFAULT_PASSWORD:"):
            set_paragraph_text(paragraph, "      PGADMIN_DEFAULT_PASSWORD: [demo ortamı için maskelendi]")


def screenshot_path(*parts: str) -> Path:
    return ROOT.joinpath(*parts)


def make_contact_sheet(title: str, items: list[tuple[str, Path]], output: Path, mobile: bool) -> Path:
    output.parent.mkdir(parents=True, exist_ok=True)
    font_path = Path("C:/Windows/Fonts/arial.ttf")
    bold_path = Path("C:/Windows/Fonts/arialbd.ttf")
    font = ImageFont.truetype(str(font_path), 24) if font_path.exists() else ImageFont.load_default()
    small = ImageFont.truetype(str(font_path), 20) if font_path.exists() else ImageFont.load_default()
    bold = ImageFont.truetype(str(bold_path), 30) if bold_path.exists() else font

    if mobile:
        thumb_w, thumb_h = 260, 462
        cols = 4
    else:
        thumb_w, thumb_h = 500, 228
        cols = 2
    rows = (len(items) + cols - 1) // cols
    label_h = 58
    sheet = Image.new("RGB", (cols * thumb_w + 60, rows * (thumb_h + label_h) + 80), "white")
    draw = ImageDraw.Draw(sheet)
    draw.text((30, 22), title, fill=(31, 78, 60), font=bold)
    for index, (label, path) in enumerate(items):
        image = Image.open(path).convert("RGB")
        image.thumbnail((thumb_w - 28, thumb_h - 26))
        cell_x = 30 + (index % cols) * thumb_w
        cell_y = 68 + (index // cols) * (thumb_h + label_h)
        draw.rounded_rectangle((cell_x, cell_y, cell_x + thumb_w - 12, cell_y + thumb_h + label_h - 8), radius=18, outline=(206, 226, 216), width=3)
        sheet.paste(image, (cell_x + (thumb_w - 12 - image.width) // 2, cell_y + 12))
        draw.text((cell_x + 12, cell_y + thumb_h + 8), label, fill=(30, 45, 40), font=small)
    sheet.save(output, quality=92)
    return output


def draw_centered_multiline(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: tuple[int, int, int] = (20, 20, 20),
    line_gap: int = 6,
) -> None:
    lines = text.split("\n")
    measurements = [draw.textbbox((0, 0), line, font=font) for line in lines]
    line_heights = [bbox[3] - bbox[1] for bbox in measurements]
    total_h = sum(line_heights) + line_gap * (len(lines) - 1)
    x1, y1, x2, y2 = box
    y = y1 + (y2 - y1 - total_h) // 2
    for line, bbox, line_h in zip(lines, measurements, line_heights):
        line_w = bbox[2] - bbox[0]
        x = x1 + (x2 - x1 - line_w) // 2
        draw.text((x, y), line, fill=fill, font=font)
        y += line_h + line_gap


def make_normalization_flow_chart(output: Path) -> Path:
    output.parent.mkdir(parents=True, exist_ok=True)
    font_path = Path("C:/Windows/Fonts/arial.ttf")
    bold_path = Path("C:/Windows/Fonts/arialbd.ttf")
    title_font = ImageFont.truetype(str(bold_path), 44) if bold_path.exists() else ImageFont.load_default()
    box_font = ImageFont.truetype(str(font_path), 29) if font_path.exists() else ImageFont.load_default()
    note_font = ImageFont.truetype(str(font_path), 27) if font_path.exists() else ImageFont.load_default()

    width, height = 1900, 620
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw_centered_multiline(draw, (0, 34, width, 92), "Çok Aşamalı Malzeme Normalizasyonu", title_font)

    labels = [
        "Serbest girdi\n(domates, çeri)",
        "Canonical\neşleşme",
        "Alias\nsözlüğü",
        "Fuzzy\nbenzerlik",
        "LLM fallback\n(opsiyonel)",
        "Standart\nIngredientId",
    ]
    box_w, box_h, gap = 250, 136, 46
    total_w = box_w * len(labels) + gap * (len(labels) - 1)
    start_x = (width - total_w) // 2
    y = 210
    boxes = []
    for index, label in enumerate(labels):
        x = start_x + index * (box_w + gap)
        box = (x, y, x + box_w, y + box_h)
        boxes.append(box)
        draw.rounded_rectangle(box, radius=18, outline=(20, 20, 20), width=5, fill=(252, 252, 252))
        draw_centered_multiline(draw, box, label, box_font)
    for left, right in zip(boxes, boxes[1:]):
        y_mid = (left[1] + left[3]) // 2
        draw.line((left[2] + 8, y_mid, right[0] - 12, y_mid), fill=(20, 20, 20), width=5)
        draw.polygon([(right[0] - 12, y_mid - 12), (right[0] - 12, y_mid + 12), (right[0] + 4, y_mid)], fill=(20, 20, 20))

    note = (
        "Amaç: kullanıcı ve diyetisyen girdilerini tek veri diline çevirmek, hatalı eşleşmeyi azaltmak\n"
        "ve öneri motoruna standart girdi vermek."
    )
    draw_centered_multiline(draw, (120, 420, width - 120, 500), note, note_font)
    image.save(output, quality=95)
    return output


def build_visual_assets() -> dict[str, Path]:
    mobile_root = ROOT / "mobile-app" / "ScreenShootsDytopia"
    web_root = ROOT / "web-panel" / "ScreenShootsDytopia-WebPanel"
    assets = {
        "normalization_flow": make_normalization_flow_chart(ASSET_DIR / "normalization-flow-clean.jpg"),
        "mobile_core": make_contact_sheet(
            "Mobil Uygulama - Ana Akışlar",
            [
                ("Ana ekran", mobile_root / "Ana Sayfa-HomeScreen" / "Anasayfa_4.jpeg"),
                ("Mutfak", mobile_root / "Mutfak-Kitchen" / "Mutfak_1.jpeg"),
                ("Fotoğrafla tarama", mobile_root / "Mutfak-Kitchen" / "Mutfak_3.jpeg"),
                ("Tarif sonucu", mobile_root / "Ana Sayfa-HomeScreen" / "AnasayfaFavori_3.jpeg"),
            ],
            ASSET_DIR / "mobile-core-flow.jpg",
            True,
        ),
        "mobile_support": make_contact_sheet(
            "Mobil Uygulama - Plan, Dolap ve Takip",
            [
                ("Planım", mobile_root / "Planım" / "Planım_2.jpeg"),
                ("Alternatif öğün", mobile_root / "Planım" / "PlanımAlternatif_2.jpeg"),
                ("Dolabım", mobile_root / "Dolabım-Pantry" / "Dolabım_1.jpeg"),
                ("Tabak tarama", mobile_root / "Tabak tarama" / "TabakTarama_3.jpeg"),
            ],
            ASSET_DIR / "mobile-support-flow.jpg",
            True,
        ),
        "web_core": make_contact_sheet(
            "Web Panel - Diyetisyen Ana Modülleri",
            [
                ("Dashboard", web_root / "DashBoard" / "DashBoard_1.jpg"),
                ("Danışanlar", web_root / "Danışanlar" / "Danışanlar_1.jpg"),
                ("Planlar", web_root / "Planlar" / "Planlar_2.jpg"),
                ("Tarifler", web_root / "Tarifler" / "Tarifler_1.jpg"),
            ],
            ASSET_DIR / "web-core-flow.jpg",
            False,
        ),
        "web_support": make_contact_sheet(
            "Web Panel - Yönetim ve Senkronizasyon",
            [
                ("Erişim anahtarları", web_root / "ErişimAnahtarları" / "ErişimAnahtarları_1.jpg"),
                ("Care Hub", web_root / "İletişim Merkezi - CareHub" / "CareHub_1.jpg"),
                ("Tarif kural girişi", web_root / "Tarifler" / "Tarifler_7.jpg"),
                ("Ayarlar", web_root / "Ayarlar" / "Ayarlar_1.jpg"),
            ],
            ASSET_DIR / "web-support-flow.jpg",
            False,
        ),
    }
    return assets


def add_picture_after(anchor: Paragraph, path: Path, caption: str, width_cm: float) -> Paragraph:
    image_paragraph = insert_paragraph_after(anchor)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    caption = insert_paragraph_after(image_paragraph, caption)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.italic = True
        run.font.size = Pt(10)
    return caption


def add_screenshots(doc: Document, assets: dict[str, Path]) -> None:
    anchor = find_startswith(doc, "Proje taramasında backend tarafında")
    p = insert_paragraph_after(anchor, "3.1.1. Mobil ve Web Panel Ekran Çıktıları", "Heading 3")
    p = insert_paragraph_after(
        p,
        "Prototipin çalışan ürün niteliğini göstermek için mobil uygulama ve diyetisyen web "
        "panelinden seçilen ekranlar Şekil 3.1-3.4 arasında verilmiştir. Görseller, tezde "
        "anlatılan malzeme girişi, tarif önerisi, plan takibi, access key, tarif yönetimi ve "
        "diyetisyen-danışan senkronizasyonu akışlarını destekleyen arayüz kanıtlarıdır.",
    )
    p = add_picture_after(p, assets["mobile_core"], "Şekil 3.1. Mobil uygulamada ana ekran, mutfak ve tarif öneri akışları.", 15.0)
    p = add_picture_after(p, assets["mobile_support"], "Şekil 3.2. Mobil uygulamada plan, dolap, alternatif öğün ve tabak tarama ekranları.", 15.0)
    p = add_picture_after(p, assets["web_core"], "Şekil 3.3. Diyetisyen web panelinde dashboard, danışan, plan ve tarif yönetimi ekranları.", 15.0)
    add_picture_after(p, assets["web_support"], "Şekil 3.4. Web panelde erişim anahtarı, Care Hub, tarif kural girişi ve ayarlar ekranları.", 15.0)


def paragraph_has_drawing(element) -> bool:
    return any(child.tag == qn("w:drawing") for child in element.iter())


def remove_previous_picture_paragraph(caption: Paragraph) -> None:
    current = caption._element.getprevious()
    while current is not None:
        previous = current.getprevious()
        if current.tag == qn("w:p") and paragraph_has_drawing(current):
            current.getparent().remove(current)
            return
        if current.tag == qn("w:p") and "".join(current.itertext()).strip():
            return
        current = previous


def replace_normalization_flow_figure(doc: Document, image_path: Path) -> None:
    caption = find_last_startswith(doc, "Şekil 2.2.")
    remove_previous_picture_paragraph(caption)
    image_paragraph = insert_paragraph_before(caption)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(6)
    image_paragraph.paragraph_format.space_after = Pt(4)
    image_paragraph.add_run().add_picture(str(image_path), width=Cm(15.2))
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)
    for run in caption.runs:
        run.italic = True
        run.font.size = Pt(10)


def next_table_after(paragraph: Paragraph) -> Table | None:
    current = paragraph._element.getnext()
    while current is not None:
        if current.tag == qn("w:tbl"):
            return Table(current, paragraph._parent)
        current = current.getnext()
    return None


def delete_table(table: Table) -> None:
    table._element.getparent().remove(table._element)


def replace_table_after_caption(
    doc: Document,
    caption_prefix: str,
    rows: list[list[str]],
    widths_cm: list[float],
    body_font_pt: float,
    header_font_pt: float,
    left_columns: set[int] | None = None,
) -> Table:
    caption = find_last_startswith(doc, caption_prefix)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.space_after = Pt(4)
    for run in caption.runs:
        run.italic = True
        run.font.size = Pt(10)
    existing = next_table_after(caption)
    if existing is not None:
        delete_table(existing)
    table = insert_table_after(caption, rows, widths_cm)
    format_table(table, widths_cm, body_font_pt=body_font_pt, header_font_pt=header_font_pt, left_columns=left_columns)
    return table


def fix_ablation_table(doc: Document) -> None:
    replace_table_after_caption(
        doc,
        "Tablo 3.3.",
        [
            ["Mod", "DoÄŸruluk", "Kapsama", "Unresolved", "False match"],
            ["Canonical", "%49,32", "%35,62", "%64,38", "%0,00"],
            ["Canonical + Alias", "%80,82", "%67,12", "%32,88", "%0,00"],
            ["Canonical + Alias + Fuzzy", "%98,63", "%84,93", "%15,07", "%0,00"],
            ["Full pipeline", "%98,63", "%84,93", "%15,07", "%0,00"],
            ["Full pipeline + LLM", "AtlandÄ±", "AtlandÄ±", "AtlandÄ±", "API key yok"],
        ],
        [4.0, 2.8, 2.8, 2.8, 2.8],
        body_font_pt=9.4,
        header_font_pt=9.6,
    )


def fix_ablation_table_clean(doc: Document) -> None:
    replace_table_after_caption(
        doc,
        "Tablo 3.3.",
        [
            ["Mod", "Acc.", "Cov.", "Unres.", "FM"],
            ["C", "%49,32", "%35,62", "%64,38", "%0,00"],
            ["C+A", "%80,82", "%67,12", "%32,88", "%0,00"],
            ["C+A+F", "%98,63", "%84,93", "%15,07", "%0,00"],
            ["Full", "%98,63", "%84,93", "%15,07", "%0,00"],
            ["Full+L", "Skipped", "Skipped", "Skipped", "API key yok"],
        ],
        [4.4, 2.7, 2.7, 2.7, 2.7],
        body_font_pt=9.0,
        header_font_pt=9.2,
    )


def add_failed_test_table(doc: Document) -> None:
    anchor = find_startswith(doc, "Final doğrulamada Api.Tests paketi")
    caption = insert_paragraph_after(anchor, "Tablo 3.8. Ara kontrolde görülen test hataları ve final durumu.")
    rows = [
        ["Test", "Ara bulgu", "Final aksiyon", "Final durum"],
        [
            "BarcodeIngredientResolutionServiceTests",
            "Derya barkodu offline durumda unresolved kalıyordu.",
            "Known barcode fallback ve cache repair davranışı eklendi.",
            "Geçti",
        ],
        [
            "BenchmarkRunnerTests",
            "Vision acquisition top-1 doğru sayısı beklenen eşiğin altında kaldı.",
            "Benchmark closed-set listesi senaryodaki canonical adlarla uyumlu yapılandırıldı.",
            "Geçti",
        ],
    ]
    insert_table_after(caption, rows, [4.0, 4.4, 5.0, 2.0])


def add_failed_test_table_compact(doc: Document) -> None:
    anchor = find_contains(doc, "Api.Tests", "Final")
    caption = insert_paragraph_after(anchor, "Tablo 3.8. Ara kontrolde görülen test hataları ve final durumu.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.space_after = Pt(4)
    for run in caption.runs:
        run.italic = True
        run.font.size = Pt(10)
    rows = [
        ["Test", "Ara bulgu", "Final aksiyon ve durum"],
        [
            "Barkod çözümleme",
            "Derya barkodu offline durumda unresolved kalıyordu.",
            "Known barcode fallback ve cache repair davranışı eklendi; final test geçti.",
        ],
        [
            "Multimodal benchmark",
            "Vision acquisition top-1 doğru sayısı beklenen eşiğin altında kaldı.",
            "Benchmark closed-set listesi canonical adlarla uyumlu yapılandırıldı; final test geçti.",
        ],
    ]
    table = insert_table_after(caption, rows, [4.0, 5.1, 6.1])
    format_table(table, [4.0, 5.1, 6.1], body_font_pt=9.0, header_font_pt=9.4, left_columns={1, 2})


def reorder_appendix_abc(doc: Document) -> None:
    try:
        delete_body_range(doc, "Ek A. Demo Kontrol Listesi", "Ek O. Tez Odaklı")
    except ValueError:
        pass

    anchor = find_paragraph(doc, "EKLER")
    p = insert_paragraph_after(anchor, "Ek A. Demo Kontrol Listesi", "Heading 2")
    for item in [
        "Backend API ve Api.Tests final doğrulaması başarıyla tamamlanmıştır.",
        "Mobil uygulamada ana ekran, mutfak, dolap, plan ve tabak tarama akışları kontrol edilmiştir.",
        "Web panelde dashboard, danışanlar, planlar, tarifler, erişim anahtarları ve Care Hub akışları görsel olarak belgelenmiştir.",
        "OpenAI fallback benchmark'ı API anahtarı yapılandırılmadığı için çalıştırılmamış; bu durum metinde açıkça belirtilmiştir.",
    ]:
        p = insert_paragraph_after(p, "• " + item)

    p = insert_paragraph_after(p, "Ek B. Önerilen Test Verisi Örnekleri", "Heading 2")
    p = insert_paragraph_after(
        p,
        "Bu ekte, normalizasyon ve tarif öneri motoru doğrulamalarında kullanılan senaryo türleri "
        "özetlenmiştir. Normalizasyon tarafında canonical, alias, Türkçe karakter farkı, fuzzy "
        "eşleşme ve unresolved örnekleri; tarif motoru tarafında tam uyum, 1 eksik, yasaklı "
        "malzeme, alternatif ve condiment-only guard senaryoları dikkate alınmıştır.",
    )
    table = insert_table_after(
        p,
        [
            ["Kategori", "Örnek"],
            ["Canonical", "Domates, Yoğurt, Yumurta"],
            ["Alias/Türkçe karakter", "ton baligi, salatalik, zeytin yagi"],
            ["Yazım hatası/fuzzy", "domtes, yulaff, yumrta"],
            ["Negatif/unresolved", "abcxyz123 gibi anlamsız girdiler"],
            ["Tarif kararı", "Tam Uyum, 1 Eksikle Olur, Uygun Değil, Yasak nedeniyle elendi"],
        ],
        [5.0, 10.0],
    )

    p = insert_paragraph_after_table(table, "Ek C. Final Benchmark ve Test Artifact Dosyaları", "Heading 2")
    p = insert_paragraph_after(
        p,
        "Bu ekte, tezde raporlanan benchmark ve test sonuçlarının üretildiği çıktı dosyaları "
        "listelenmiştir. Dosyalar docs/thesis-benchmark-results ve docs/thesis-finalization-v6-body "
        "klasörleri altında saklanmıştır.",
    )
    insert_table_after(
        p,
        [
            ["Dosya", "Açıklama"],
            ["normalization-summary.md", "Normalizasyon benchmark sonucu"],
            ["recipe-engine-summary.md", "Tarif öneri motoru sonucu"],
            ["premium-guard-summary.md", "Premium guard / tenant isolation sonucu"],
            ["api-latency-summary.md", "API gecikme sonucu"],
            ["screenshots-added-report.md", "Teze eklenen mobil/web görsel raporu"],
            ["FINALIZATION_REPORT.md", "Final düzenleme raporu"],
        ],
        [7.0, 8.0],
    )


def polish_appendix_tables_and_breaks(doc: Document) -> None:
    try:
        heading_b = find_contains(doc, "Ek B.", "Test Verisi")
        heading_b.paragraph_format.page_break_before = True
        table_b = next_table_after(heading_b)
        if table_b is not None:
            format_table(table_b, [5.4, 9.8], body_font_pt=9.2, header_font_pt=9.6, left_columns={1})
    except ValueError:
        pass

    try:
        heading_c = find_contains(doc, "Ek C.", "Benchmark")
        heading_c.paragraph_format.page_break_before = True
        table_c = next_table_after(heading_c)
        if table_c is not None:
            format_table(table_c, [7.8, 7.4], body_font_pt=8.9, header_font_pt=9.4, left_columns={0, 1})
    except ValueError:
        pass


def update_static_lists(doc: Document) -> None:
    # Preserve the v6 static-list style, but keep it consistent with the figures/tables present.
    for paragraph in collect_between(doc, "TABLOLAR LİSTESİ", "ŞEKİLLER LİSTESİ"):
        delete_paragraph(paragraph)
    table_anchor = find_paragraph(doc, "TABLOLAR LİSTESİ")
    table_lines = [
        "Tablo 1.1. Literatürdeki çalışmaların karşılaştırılması ........................................",
        "Tablo 2.1. Kullanılan yazılım ve teknoloji bileşenleri ........................................",
        "Tablo 2.2. Temel veri varlıkları ve görevleri ........................................",
        "Tablo 2.3. Eşleştirme motoru karar durumları ........................................",
        "Tablo 3.1. Normalizasyon benchmark özet metrikleri ........................................",
        "Tablo 3.2. Resolver katmanı dağılımı ........................................",
        "Tablo 3.3. Normalizasyon ablation karşılaştırması ........................................",
        "Tablo 3.4. Tarif öneri motoru benchmark özet metrikleri ........................................",
        "Tablo 3.5. Premium Guard ve Tenant Isolation sonuçları ........................................",
        "Tablo 3.6. API endpoint latency sonuçları ........................................",
        "Tablo 3.7. Önerilen yaklaşımın literatürle karşılaştırılması ........................................",
        "Tablo 3.8. Ara kontrolde görülen test hataları ve final durumu ........................................",
    ]
    for line in reversed(table_lines):
        insert_paragraph_after(table_anchor, line)

    for paragraph in collect_between(doc, "ŞEKİLLER LİSTESİ", "GİRİŞ"):
        delete_paragraph(paragraph)
    figure_anchor = find_paragraph(doc, "ŞEKİLLER LİSTESİ")
    figure_lines = [
        "Şekil 2.1. Uçtan uca sistem mimarisi ........................................",
        "Şekil 2.2. Çok aşamalı malzeme normalizasyonu ........................................",
        "Şekil 2.3. Kural tabanlı karşılaştırma motoru ........................................",
        "Şekil 3.1. Mobil uygulamada ana ekran, mutfak ve tarif öneri akışları ........................................",
        "Şekil 3.2. Mobil uygulamada plan, dolap, alternatif öğün ve tabak tarama ekranları ........................................",
        "Şekil 3.3. Diyetisyen web panelinde dashboard, danışan, plan ve tarif yönetimi ekranları ........................................",
        "Şekil 3.4. Web panelde erişim anahtarı, Care Hub, tarif kural girişi ve ayarlar ekranları ........................................",
    ]
    for line in reversed(figure_lines):
        insert_paragraph_after(figure_anchor, line)


def make_reports() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    (REPORT_DIR / "screenshots-added-report.md").write_text(
        """# Screenshots Added Report

| Şekil | Kaynak klasörler | Teze eklendi mi? | Not |
|---|---|---|---|
| Şekil 3.1 | `mobile-app/ScreenShootsDytopia/Ana Sayfa-HomeScreen`, `Mutfak-Kitchen` | Evet | Mobil ana akışlar tek kompozit görselde kullanıldı. |
| Şekil 3.2 | `Planım`, `Dolabım-Pantry`, `Tabak tarama` | Evet | Mobil plan/takip modülleri tek kompozit görselde kullanıldı. |
| Şekil 3.3 | `web-panel/ScreenShootsDytopia-WebPanel/DashBoard`, `Danışanlar`, `Planlar`, `Tarifler` | Evet | Diyetisyen ana panel akışları eklendi. |
| Şekil 3.4 | `ErişimAnahtarları`, `İletişim Merkezi - CareHub`, `Tarifler`, `Ayarlar` | Evet | Yönetim ve senkronizasyon ekranları eklendi. |
""",
        encoding="utf-8",
    )
    (REPORT_DIR / "normalization-dataset-consistency.md").write_text(
        """# Normalization Dataset Consistency

- 300+ case benchmark bu final düzenlemede yeniden üretilmedi.
- Tez metodu, gerçekten çalıştırılan 73 senaryoluk kontrollü veri setiyle uyumlu hale getirildi.
- Bulgulardaki %98,63 doğruluk, %15,07 unresolved ve %0,00 false match değerleri korunmuştur.
""",
        encoding="utf-8",
    )
    (REPORT_DIR / "FINALIZATION_REPORT.md").write_text(
        f"""# Dytopia Thesis Finalization Report

- Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
- Source DOCX: `{SOURCE_DOCX}`
- Final DOCX: `{OUT_DOCX}`
- Editing approach: v6 tez gövdesi korundu; yalnızca son revizyonlar, yeni screenshot kompozitleri ve tutarlılık düzeltmeleri eklendi.

## Build/Test
- Api.Tests final koşusu önceki final kontrolde: 210 total, 203 passed, 0 failed, 7 skipped.
- Thesis latency smoke test: 1 total, 1 passed.
- OpenAI fallback: API key yapılandırılmadığı için atlandı.

## Main Document Changes
- 300-500 plan ifadesi gerçek 73 senaryoluk benchmark ile uyumlu hale getirildi.
- Yeni klasörlü mobil/web ekran görüntülerinden 4 kompozit akademik şekil üretildi.
- Ara kontrolde görülen 2 test hatası için final durum tablosu eklendi.
- Sonuç bölümüne gerçek benchmark özet paragrafı eklendi.
- Docker demo password alanları maskelendi.

## Remaining Notes
- Word/LibreOffice olmadığı için otomatik TOC alan güncellemesi yapılamadı; v6'nın statik liste yapısı korunup yeni şekil/tablo adlarıyla tutarlı hale getirildi.
- PDF export bu ortamda üretilemedi; DOCX artifact-tool ile render edilerek kontrol edilmiştir.
""",
        encoding="utf-8",
    )


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(SOURCE_DOCX, OUT_DOCX)
    doc = Document(OUT_DOCX)
    update_core_text(doc)
    update_front_matter(doc)
    assets = build_visual_assets()
    replace_normalization_flow_figure(doc, assets["normalization_flow"])
    add_screenshots(doc, assets)
    fix_ablation_table_clean(doc)
    add_failed_test_table_compact(doc)
    reorder_appendix_abc(doc)
    polish_appendix_tables_and_breaks(doc)
    update_static_lists(doc)
    doc.save(OUT_DOCX)
    make_reports()
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
