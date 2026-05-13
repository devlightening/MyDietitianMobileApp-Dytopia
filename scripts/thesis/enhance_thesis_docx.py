from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.shared import Cm
from docx.text.paragraph import Paragraph

ROOT = Path(r"C:/Users/hy971/source/repos/MyDietitianMobileApp")
IN_PATH = Path(r"C:/Users/hy971/Desktop/docs/thesis/Dytopia_Bitirme_Tezi_Profesyonel_Taslak.docx")
OUT_PATH = ROOT / "tmp-build" / "Dytopia_Bitirme_Tezi_Profesyonel_Taslak_v2.docx"

AUTHOR = "Halil İbrahim YILDIRIM"
STUDENT_NO = "212503009"
TITLE = (
    "SERBEST METİN MALZEME GİRDİLERİNİN ÇOK AŞAMALI NORMALİZASYONU VE TAKSONOMİ TABANLI "
    "AÇIKLANABİLİR TARİF ÖNERİ SİSTEMİ: DİYETİSYEN DESTEKLİ BİR PLATFORM"
)


def insert_paragraph_after(paragraph, text: str | None = None, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def read_controller_names() -> list[str]:
    p = ROOT / "src/MyDietitianMobileApp.Api/Controllers"
    return sorted([x.stem for x in p.glob("*.cs")])


def read_mobile_screens() -> list[str]:
    p = ROOT / "mobile-app/src/screens"
    return sorted([x.stem for x in p.glob("*.tsx")])


def read_web_routes() -> list[str]:
    p = ROOT / "web-panel/app/dashboard"
    routes = []
    for x in p.rglob("page.tsx"):
        routes.append(str(x.relative_to(ROOT / "web-panel/app")).replace('\\', '/'))
    return sorted(routes)


def read_dbsets() -> list[str]:
    appdb = ROOT / "src/MyDietitianMobileApp.Infrastructure/Persistence/AppDbContext.cs"
    lines = appdb.read_text(encoding="utf-8", errors="ignore").splitlines()
    out = []
    for ln in lines:
        if "DbSet<" in ln:
            m = re.search(r"DbSet<([^>]+)>\s+([^\s{]+)", ln)
            if m:
                out.append(f"{m.group(1)} → {m.group(2)}")
    return out


def read_domain_entities() -> list[str]:
    p = ROOT / "src/MyDietitianMobileApp.Domain/Entities"
    return sorted([x.stem for x in p.glob("*.cs")])


def read_test_files() -> list[str]:
    p = ROOT / "tests/MyDietitianMobileApp.Api.Tests"
    files = sorted([str(x.relative_to(ROOT / "tests")) for x in p.rglob("*.cs")])
    # keep it reasonable but still long enough to add pages
    return files


def read_mobile_api_modules() -> list[str]:
    p = ROOT / "mobile-app/src/api"
    return sorted([x.name for x in p.glob("*.ts")])


def read_web_api_modules() -> list[str]:
    p = ROOT / "web-panel/lib/api"
    return sorted([x.name for x in p.glob("*.ts")])


def read_program_service_registrations() -> list[str]:
    program = ROOT / "src/MyDietitianMobileApp.Api/Program.cs"
    if not program.exists():
        return []
    lines = program.read_text(encoding="utf-8", errors="ignore").splitlines()
    out: list[str] = []
    for ln in lines:
        if "builder.Services.Add" in ln:
            cleaned = ln.strip().rstrip(";")
            # keep short
            out.append(cleaned)
    return out


def read_docker_compose_excerpt() -> list[str]:
    p = ROOT / "docker-compose.yml"
    if not p.exists():
        return []
    lines = p.read_text(encoding="utf-8", errors="ignore").splitlines()
    # Keep whole file; usually short.
    return lines[:200]


def read_api_endpoint_inventory() -> list[str]:
    """Extract a best-effort endpoint inventory from controller attributes."""
    controllers_dir = ROOT / "src/MyDietitianMobileApp.Api/Controllers"
    out: list[str] = []
    http_attr = re.compile(r"\[(HttpGet|HttpPost|HttpPut|HttpDelete|HttpPatch)(?:\(([^)]*)\))?\]")
    method_sig = re.compile(
        r"\b(public|private|protected)\s+(?:async\s+)?(?:Task|ActionResult|IActionResult|Task<[^>]+>|ActionResult<[^>]+>)\s+(\w+)\s*\("
    )

    for f in sorted(controllers_dir.glob("*.cs")):
        text = f.read_text(encoding="utf-8", errors="ignore").splitlines()
        controller = f.stem.replace("Controller", "")
        pending_attrs: list[str] = []
        for ln in text:
            ln_s = ln.strip()
            m = http_attr.search(ln_s)
            if m:
                verb = m.group(1)
                template = (m.group(2) or "").strip()
                pending_attrs.append(f"{verb}{'(' + template + ')' if template else ''}")
                continue
            ms = method_sig.search(ln_s)
            if ms and pending_attrs:
                method = ms.group(2)
                for a in pending_attrs:
                    out.append(f"{controller}.{method} — {a}")
                pending_attrs = []
        if not out or (out and out[-1] != ""):
            pass
    return out


def read_versions() -> dict[str, str]:
    mobile_pkg = (ROOT / "mobile-app/package.json").read_text(encoding="utf-8")
    web_pkg = (ROOT / "web-panel/package.json").read_text(encoding="utf-8")

    def find(json_text: str, key: str) -> str | None:
        m = re.search(rf"\"{re.escape(key)}\"\s*:\s*\"([^\"]+)\"", json_text)
        return m.group(1) if m else None

    out: dict[str, str] = {}
    for k in ["expo", "react", "react-native", "typescript", "axios"]:
        v = find(mobile_pkg, k)
        if v:
            out[f"mobile:{k}"] = v
    for k in ["next", "react", "typescript", "tailwindcss", "axios"]:
        v = find(web_pkg, k)
        if v:
            out[f"web:{k}"] = v

    api_csproj = next((ROOT / "src/MyDietitianMobileApp.Api").glob("*.csproj"))
    txt = api_csproj.read_text(encoding="utf-8")
    m = re.search(r"<TargetFramework>([^<]+)</TargetFramework>", txt)
    if m:
        out["backend:TargetFramework"] = m.group(1)

    return out


def find_paragraph(doc: Document, exact_text: str):
    want = exact_text.strip()
    for p in doc.paragraphs:
        if p.text.strip() == want:
            return p
    return None


def add_heading_block(anchor: Paragraph, heading: str, paras: list[str]) -> Paragraph:
    h = insert_paragraph_after(anchor, heading, style="Heading 2")
    cur = h
    for t in paras:
        cur = insert_paragraph_after(cur, t, style="Normal")
    return cur


def add_bullets(anchor: Paragraph, items: list[str]) -> Paragraph:
    cur = anchor
    for it in items:
        cur = insert_paragraph_after(cur, it, style="List Bullet")
    return cur


def main() -> None:
    if not IN_PATH.exists():
        raise SystemExit(f"Missing input docx: {IN_PATH}")

    doc = Document(str(IN_PATH))

    # Enforce margins per thesis guide (typical: left 3cm, right 2.5cm, top/bottom 2.5cm)
    for sec in doc.sections:
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.5)
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)

    # Cover title updates (first pages)
    for p in doc.paragraphs[:40]:
        txt = p.text.strip()
        if txt.isupper() and ("AKILLI MUTFAK" in txt or "DİYETİSYEN" in txt) and "TEZ" not in txt:
            p.text = TITLE

    # Insert school number after author name on cover pages (avoid duplicates)
    inserted = 0
    okul_line = f"OKUL NO: {STUDENT_NO}"
    for p in list(doc.paragraphs[:120]):
        if p.text.strip() != AUTHOR:
            continue
        next_p = p._p.getnext()
        if next_p is not None:
            try:
                next_para = Paragraph(next_p, p._parent)
                if next_para.text.strip() == okul_line:
                    continue
            except Exception:
                pass
        np = insert_paragraph_after(p, okul_line)
        np.style = p.style
        np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        inserted += 1
        if inserted >= 2:
            break

    # If the template already had school-no lines duplicated, collapse consecutive duplicates.
    for p in list(doc.paragraphs[:150]):
        if p.text.strip() != okul_line:
            continue
        prev = p._p.getprevious()
        if prev is None:
            continue
        try:
            prev_para = Paragraph(prev, p._parent)
            if prev_para.text.strip() == okul_line:
                p._p.getparent().remove(p._p)
        except Exception:
            continue

    # Expand tech stack under 2.1
    anchor = find_paragraph(doc, "2.1. Kullanılan Yazılım ve Donanım Ortamı")
    if anchor is not None:
        versions = read_versions()
        paras = [
            f"Backend: {versions.get('backend:TargetFramework','net8.0')} hedef çerçevesi üzerinde ASP.NET Core Web API mimarisi ile geliştirilmiştir. Veri erişim katmanında Entity Framework Core ve PostgreSQL sağlayıcısı (Npgsql) kullanılmıştır.",
            f"Mobil Uygulama: Expo ({versions.get('mobile:expo','-')}), React ({versions.get('mobile:react','-')}), React Native ({versions.get('mobile:react-native','-')}) bileşenleri ile geliştirilmiş; veri erişiminde Axios ({versions.get('mobile:axios','-')}) ve durum yönetiminde React Query yaklaşımı tercih edilmiştir.",
            f"Web Panel: Next.js ({versions.get('web:next','-')}) üzerinde React ({versions.get('web:react','-')}) ile geliştirilmiş; arayüz bileşenleri ve stil yönetimi için Tailwind CSS ({versions.get('web:tailwindcss','-')}) kullanılmıştır.",
            "Gerçek zamanlı bildirim/aktivite akışları için SignalR temelli iletişim katmanı desteklenmiştir.",
        ]
        add_heading_block(anchor, "2.1.1. Teknoloji Yığını ve Sürüm Bilgileri", paras)

    # Add a concrete module inventory section after 2.9 (fits the \"hoca ne varsa yazın\" request)
    anchor_29 = find_paragraph(doc, "2.9. Test ve Değerlendirme Yöntemi")
    if anchor_29 is not None:
        paras = [
            "Bu tez kapsamında geliştirilen sistem; mobil uygulama (danışan), web panel (diyetisyen) ve backend API (iş kuralları + veri) katmanlarından oluşmaktadır. Proje taraması, kaynak kod klasörleri ve API denetleyicileri incelenerek yapılmıştır.",
            "Backend API tarafında; kimlik doğrulama (JWT), yetkilendirme, müşteri/diyetisyen yönetimi, tarif yönetimi, mutfak (kitchen) eşleştirme, alışveriş listesi, pantry yönetimi, gamification (oyunlaştırma), bildirim tercihleri, ölçüm takibi, barcode eşleştirme ve benchmark/diagnostics uç noktaları yer almaktadır.",
            "Mobil uygulamada; giriş/kayıt, dashboard, profil, ölçüm takibi, su takibi, pantry, alışveriş listesi, mutfak sonuç ekranı, tarif detayları, barkod tarama, fiş tarama ve oyunlaştırma ekranları gibi modüller bulunmaktadır.",
            "Web panelde; dashboard, danışan yönetimi, randevular, tarif oluşturma/düzenleme, tarif içe aktarma, recipe-match/simülasyon, erişim anahtarı (access key) yönetimi, branding ve ayarlar sayfaları bulunmaktadır.",
        ]
        add_heading_block(anchor_29, "2.10. Katmanlar ve Modül Envanteri (Proje Taraması)", paras)

    # Fill 2.9 (currently empty in template): test & evaluation methodology
    h_29 = find_paragraph(doc, "2.9. Test ve Değerlendirme Yöntemi")
    if h_29 is not None:
        # Only insert if still empty (next paragraph is a heading)
        nxt = h_29._p.getnext()
        is_empty = True
        if nxt is not None:
            try:
                nxt_para = Paragraph(nxt, h_29._parent)
                if not nxt_para.style.name.startswith("Heading") and nxt_para.text.strip():
                    is_empty = False
            except Exception:
                pass
        if is_empty:
            cur = insert_paragraph_after(
                h_29,
                "Bu çalışmada değerlendirme, yalnızca yazılım doğrulama (QA) düzeyinde bırakılmayıp; normalizasyon hattı ve kural tabanlı karar motorunun ölçülebilir performansını gösterecek biçimde tasarlanmıştır.",
                style="Normal",
            )
            cur = insert_paragraph_after(cur, "2.9.1. Değerlendirme Veri Setleri", style="Heading 3")
            cur = insert_paragraph_after(
                cur,
                "Normalizasyon için altın standart bir veri seti tanımlanır. Bu veri seti; gerçekçi serbest metin malzeme ifadeleri (yazım hataları, Türkçe karakter eksikleri, kısaltmalar, eşanlamlılar) ile bunların beklenen canonical kimlik eşleşmelerini içerir. Hedef büyüklük 300–500 ifade olacak şekilde planlanmıştır.",
                style="Normal",
            )
            cur = insert_paragraph_after(
                cur,
                "Karar motoru için ayrı bir senaryo veri seti tanımlanır. Bu sette; “sepet (kullanıcı malzemeleri) + tarif kural rolleri (zorunlu/opsiyonel/yasak/alternatif)” bileşenleri ve beklenen karar sınıfı (Tam Uyum / 1 Eksikle Olur / Uygun Değil) bulunur. Senaryolar, yasaklı malzeme çakışmaları ve alternatif kullanımı gibi zor durumları içerecek şekilde dengelenir.",
                style="Normal",
            )
            cur = insert_paragraph_after(cur, "2.9.2. Baseline ve Ablation Tasarımı", style="Heading 3")
            cur = insert_paragraph_after(
                cur,
                "Normalizasyon hattı için ablation karşılaştırması yapılır: (i) yalnız canonical, (ii) canonical+alias, (iii) canonical+alias+fuzzy, (iv) tam hat (opsiyonel LLM fallback dahil). Bu sayede her katmanın kapsama/doğruluk üzerindeki katkısı görünür kılınır.",
                style="Normal",
            )
            cur = insert_paragraph_after(
                cur,
                "Karar motoru için de benzer şekilde karşılaştırma yapılır: (i) taksonomi kuralları devre dışı, (ii) alternatif/ikame kuralları devre dışı, (iii) tam sistem. Böylece “taksonomi tabanlı” iddiası nicel olarak desteklenebilir.",
                style="Normal",
            )
            cur = insert_paragraph_after(cur, "2.9.3. Metrikler ve Ölçüm Protokolü", style="Heading 3")
            cur = insert_paragraph_after(
                cur,
                "Normalizasyon için temel metrikler: doğruluk (accuracy), top‑k başarı, kapsama (hangi katmanda eşleşme bulunduğu) ve gecikme (latency) ölçümleridir. Karar motoru için sınıflandırma doğruluğu, yanlış pozitif/negatif oranları ve açıklama alanlarının doluluk oranı raporlanır.",
                style="Normal",
            )
            cur = insert_paragraph_after(
                cur,
                "Süre ölçümleri için her test senaryosu birden fazla tekrar çalıştırılarak ortalama/medyan ve p95 değerleri raporlanır. Ölçümlerde aynı makine ve aynı sürüm (aynı commit) kullanılır; sonuçlar Tablo 3.x altında sunulur.",
                style="Normal",
            )
            cur = insert_paragraph_after(cur, "2.9.4. Geçerlik Tehditleri", style="Heading 3")
            cur = insert_paragraph_after(
                cur,
                "Veri setinin temsil gücü, alias sözlüğünün kapsamı ve fuzzy eşik değerleri sonuçları etkileyebilir. Bu nedenle veri seti dağılımı (canonical/alias/fuzzy/LLM) kategorilerine göre dengelenir ve hatalar tür bazında raporlanır. LLM fallback kullanılıyorsa, karar verici rolü sınırlanır ve yalnızca belirsiz girdilerde öneri amaçlı kullanıldığı açıkça belirtilir.",
                style="Normal",
            )

    # Appendices: inventory extracted from repo
    ekler = find_paragraph(doc, "EKLER")
    if ekler is not None:
        tail = add_heading_block(ekler, "Ek D. Backend API Uç Noktaları (Denetleyici Listesi)", [
            "Bu ekte, projede yer alan backend API denetleyicileri (controller) listelenmiştir. Denetleyici isimleri proje kaynak kodundan otomatik çıkarılmıştır."
        ])
        tail = add_bullets(tail, read_controller_names())

        tail = add_heading_block(tail, "Ek E. Mobil Uygulama Ekranları", [
            "Bu ekte, mobil uygulamada yer alan ekranlar (screen) listelenmiştir. Liste proje kaynak kodundan otomatik çıkarılmıştır."
        ])
        tail = add_bullets(tail, read_mobile_screens())

        tail = add_heading_block(tail, "Ek F. Web Panel Dashboard Sayfaları", [
            "Bu ekte, web panel dashboard altında bulunan sayfalar listelenmiştir. Liste proje kaynak kodundan otomatik çıkarılmıştır."
        ])
        tail = add_bullets(tail, read_web_routes())

        tail = add_heading_block(tail, "Ek G. Veritabanı Şeması (DbSet Envanteri)", [
            "Bu ekte, AppDbContext üzerinde tanımlı DbSet'ler listelenmiştir. Bu liste veritabanındaki temel tabloları/ilişkileri yansıtan üst seviye envanter niteliğindedir."
        ])
        tail = add_bullets(tail, read_dbsets())

        tail = add_heading_block(tail, "Ek H. Domain Entity Sınıfları", [
            "Bu ekte, Domain katmanındaki entity sınıfları listelenmiştir (kaynak koddan otomatik çıkarılmıştır)."
        ])
        tail = add_bullets(tail, read_domain_entities())

        tail = add_heading_block(tail, "Ek I. Test Envanteri (Api.Tests)", [
            "Bu ekte, test projesi altındaki test dosyaları listelenmiştir. Bulgular bölümünde sunulan metrikler için ilgili test/benchmark sınıfları referans alınmalıdır."
        ])
        tail = add_bullets(tail, read_test_files())

        tail = add_heading_block(tail, "Ek J. Mobil API Modülleri", [
            "Bu ekte, mobil uygulamanın backend ile haberleşmesinde kullanılan API modülleri (TypeScript) listelenmiştir."
        ])
        tail = add_bullets(tail, read_mobile_api_modules())

        tail = add_heading_block(tail, "Ek K. Web Panel API Modülleri", [
            "Bu ekte, web panelin backend ile haberleşmesinde kullanılan API modülleri (TypeScript) listelenmiştir."
        ])
        tail = add_bullets(tail, read_web_api_modules())

        tail = add_heading_block(tail, "Ek L. API Endpoint Envanteri (HTTP Metodu + Metot Adı)", [
            "Bu ekte, backend API denetleyici sınıflarında kullanılan HTTP attribute'larından çıkarılan uç nokta envanteri yer almaktadır. Liste otomatik çıkarım olduğu için, nihai doğrulama Swagger/OpenAPI çıktısı üzerinden yapılmalıdır."
        ])
        tail = add_bullets(tail, read_api_endpoint_inventory())

        tail = add_heading_block(tail, "Ek M. Backend Servis Kayıtları (Program.cs Özet)", [
            "Bu ekte, backend uygulamasında DI (Dependency Injection) konteynerine eklenen servis kayıtları listelenmiştir (Program.cs)."
        ])
        tail = add_bullets(tail, read_program_service_registrations())

        compose = read_docker_compose_excerpt()
        if compose:
            tail = add_heading_block(tail, "Ek N. docker-compose.yml (Altyapı Özeti)", [
                "Bu ekte, yerel geliştirme ortamı için kullanılan docker-compose yapılandırması verilmiştir."
            ])
            # Add as a preformatted-like block using Normal paragraphs (Word will keep it readable).
            for ln in compose:
                tail = insert_paragraph_after(tail, ln, style="Normal")

    # Heading font normalization (keep consistent)
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            for r in p.runs:
                r.font.name = "Times New Roman"
                if p.style.name == "Heading 1":
                    r.font.size = Pt(14)
                elif p.style.name == "Heading 2":
                    r.font.size = Pt(13)

    doc.save(str(OUT_PATH))
    print(str(OUT_PATH))


if __name__ == "__main__":
    main()
