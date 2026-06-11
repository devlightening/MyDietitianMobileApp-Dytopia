#!/usr/bin/env python3
"""
format_thesis_final.py – Dytopia Thesis DOCX Formatter
Formats Dytopia_Bitirme_Tezi_TarsusUni_Final_Akademik_Grafikli.docx
Reference format: bitirme_tezi_saim_gezer_V2.docx (style only, no content)
Output: Dytopia_Bitirme_Tezi_TarsusUni_Final_Duzenlenmis.docx
"""

import shutil
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

ROOT = Path(r'c:\Users\hy971\source\repos\MyDietitianMobileApp')
SOURCE = ROOT / 'Dytopia_Bitirme_Tezi_TarsusUni_Final_Akademik_Grafikli.docx'
BACKUP = ROOT / 'Dytopia_Bitirme_Tezi_TarsusUni_Final_Akademik_Grafikli_BACKUP.docx'
OUTPUT = ROOT / 'Dytopia_Bitirme_Tezi_TarsusUni_Final_Duzenlenmis.docx'

FONT = 'Times New Roman'
SZ = '24'  # 12pt = 24 half-points

FORBIDDEN = ['DTHybrit', 'TarsusFer', 'Saim Gezer', 'MovieLens', 'SASRec', 'GRU4Rec', 'DeepFM']

ABBREVIATIONS = [
    ('API', 'Application Programming Interface (Uygulama Programlama Arayüzü)'),
    ('B2B', 'Business to Business (İşletmeden İşletmeye)'),
    ('B2B2C', 'Business to Business to Consumer (İşletmeden Tüketiciye)'),
    ('CSV', 'Comma Separated Values (Virgülle Ayrılmış Değerler)'),
    ('EF Core', 'Entity Framework Core'),
    ('JWT', 'JSON Web Token'),
    ('LLM', 'Large Language Model (Büyük Dil Modeli)'),
    ('MVP', 'Minimum Viable Product (Minimum Uygulanabilir Ürün)'),
    ('NRS', 'Nutrition Recommendation System (Beslenme Öneri Sistemi)'),
    ('RBAC', 'Role-Based Access Control (Rol Tabanlı Erişim Kontrolü)'),
    ('REST', 'Representational State Transfer'),
    ('SaaS', 'Software as a Service (Hizmet Olarak Yazılım)'),
    ('UI', 'User Interface (Kullanıcı Arayüzü)'),
    ('UX', 'User Experience (Kullanıcı Deneyimi)'),
]

TECH_TERMS = [
    'Dytopia', 'MyDietitian', 'NutritionResolver', 'FuzzyMatcher',
    'EF Core', '.NET', 'ASP.NET', 'React Native', 'Expo',
    'PostgreSQL', 'JWT', 'RBAC', 'REST', 'API', 'LLM',
    'OpenAI', 'GPT', 'B2B', 'B2B2C', 'SaaS', 'MVP',
    'DbSet', 'LINQ', 'Controller', 'endpoint', 'backend',
    'frontend', 'middleware', 'NormalizationService',
    'RecipeEngine', 'PremiumGuard', 'AccessKey',
]

# ── Helpers ────────────────────────────────────────────────────

def _set_run_fmt(r_el):
    """Set a w:r element to TNR 12pt black, no underline."""
    rPr = r_el.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r_el.insert(0, rPr)
    # Font – replace entirely (removes theme refs)
    for old in rPr.findall(qn('w:rFonts')):
        rPr.remove(old)
    rf = OxmlElement('w:rFonts')
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(a), FONT)
    rPr.insert(0, rf)
    # Size
    for tag in ('w:sz', 'w:szCs'):
        for old in rPr.findall(qn(tag)):
            rPr.remove(old)
        el = OxmlElement(tag); el.set(qn('w:val'), SZ); rPr.append(el)
    # Color black
    for old in rPr.findall(qn('w:color')):
        rPr.remove(old)
    c = OxmlElement('w:color'); c.set(qn('w:val'), '000000'); rPr.append(c)
    # Remove underline
    for old in rPr.findall(qn('w:u')):
        rPr.remove(old)


def _clear_p(p_el):
    for ch in list(p_el):
        if ch.tag != qn('w:pPr'):
            p_el.remove(ch)


def _center(p_el):
    pPr = p_el.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr'); p_el.insert(0, pPr)
    for old in pPr.findall(qn('w:jc')):
        pPr.remove(old)
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); pPr.append(jc)


def _add_run(p_el, text, bold=False):
    r = OxmlElement('w:r'); _set_run_fmt(r)
    if bold:
        rPr = r.find(qn('w:rPr'))
        rPr.append(OxmlElement('w:b')); rPr.append(OxmlElement('w:bCs'))
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text; r.append(t); p_el.append(r)
    return r


def _page_field_runs():
    runs = []
    for ft in ('begin', 'instr', 'sep', 'disp', 'end'):
        r = OxmlElement('w:r'); _set_run_fmt(r)
        if ft == 'begin':
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin'); r.append(fc)
        elif ft == 'instr':
            it = OxmlElement('w:instrText')
            it.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            it.text = ' PAGE '; r.append(it)
        elif ft == 'sep':
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'separate'); r.append(fc)
        elif ft == 'disp':
            t = OxmlElement('w:t'); t.text = '1'; r.append(t)
        elif ft == 'end':
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'end'); r.append(fc)
        runs.append(r)
    return runs


def _insert_after(anchor, text='', bold=False):
    p = OxmlElement('w:p'); anchor.addnext(p)
    if text:
        _add_run(p, text, bold)
    return p


def _del(el):
    par = el.getparent()
    if par is not None:
        par.remove(el)


def _find_idx(doc, txt):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == txt:
            return i
    return None


def _next_h1(doc, after):
    for i, p in enumerate(doc.paragraphs):
        if i > after and p.style.name.startswith('Heading 1') and p.text.strip():
            return i
    return None


def _del_between(doc, lo, hi):
    for j in range(hi - 1, lo, -1):
        ps = doc.paragraphs
        if j < len(ps):
            _del(ps[j]._element)


# ── Steps ──────────────────────────────────────────────────────

def s01_backup():
    print('[01] Backup...')
    shutil.copy2(SOURCE, BACKUP)
    print(f'  → {BACKUP.name}')


def s02_doc_defaults(doc):
    print('[02] Document defaults & theme...')
    # docDefaults
    stEl = doc.styles.element
    dd = stEl.find(qn('w:docDefaults'))
    if dd is not None:
        rPrD = dd.find(f'{qn("w:rPrDefault")}/{qn("w:rPr")}')
        if rPrD is not None:
            for old in rPrD.findall(qn('w:rFonts')):
                rPrD.remove(old)
            rf = OxmlElement('w:rFonts')
            for a in ('w:ascii','w:hAnsi','w:cs','w:eastAsia'):
                rf.set(qn(a), FONT)
            rPrD.insert(0, rf)
            for tag in ('w:sz','w:szCs'):
                for old in rPrD.findall(qn(tag)):
                    rPrD.remove(old)
                el = OxmlElement(tag); el.set(qn('w:val'), SZ); rPrD.append(el)
    # Theme
    try:
        a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'
        for rel in doc.part.rels.values():
            if 'theme' in rel.reltype:
                te = rel.target_part.element
                for tag in ('majorFont', 'minorFont'):
                    for fs in te.iter(f'{{{a_ns}}}{tag}'):
                        lat = fs.find(f'{{{a_ns}}}latin')
                        if lat is not None:
                            lat.set('typeface', FONT)
                break
    except Exception:
        pass
    print('  → done')


def s03_styles(doc):
    print('[03] Styles...')
    n = 0
    for st in doc.styles:
        if not hasattr(st, 'font') or st.font is None:
            continue
        el = st.element
        rPr = el.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr'); el.append(rPr)
        for old in rPr.findall(qn('w:rFonts')):
            rPr.remove(old)
        rf = OxmlElement('w:rFonts')
        for a in ('w:ascii','w:hAnsi','w:cs','w:eastAsia'):
            rf.set(qn(a), FONT)
        rPr.insert(0, rf)
        for tag in ('w:sz','w:szCs'):
            for old in rPr.findall(qn(tag)):
                rPr.remove(old)
            s = OxmlElement(tag); s.set(qn('w:val'), SZ); rPr.append(s)
        for old in rPr.findall(qn('w:color')):
            rPr.remove(old)
        c = OxmlElement('w:color'); c.set(qn('w:val'), '000000'); rPr.append(c)
        for old in rPr.findall(qn('w:u')):
            rPr.remove(old)
        for old in rPr.findall(qn('w:i')):
            rPr.remove(old)
        for old in rPr.findall(qn('w:iCs')):
            rPr.remove(old)
        n += 1
    print(f'  → {n} styles')


def s04_runs(doc):
    print('[04] Runs...')
    body = doc.element.body
    n = 0
    for r_el in body.iter(qn('w:r')):
        if r_el.find(qn('w:drawing')) is not None:
            continue
        _set_run_fmt(r_el)
        n += 1
    # paragraph-level rPr
    for p_el in body.iter(qn('w:p')):
        pPr = p_el.find(qn('w:pPr'))
        if pPr is None:
            continue
        rPr = pPr.find(qn('w:rPr'))
        if rPr is None:
            continue
        for old in rPr.findall(qn('w:rFonts')):
            for a in ('w:ascii','w:hAnsi','w:cs','w:eastAsia'):
                old.set(qn(a), FONT)
            for ta in ('w:asciiTheme','w:hAnsiTheme','w:cstheme','w:eastAsiaTheme'):
                if qn(ta) in old.attrib:
                    del old.attrib[qn(ta)]
        for tag in ('w:sz','w:szCs'):
            for el in rPr.findall(qn(tag)):
                el.set(qn('w:val'), SZ)
        for el in rPr.findall(qn('w:color')):
            el.set(qn('w:val'), '000000')
            for at in [qn('w:themeColor'), qn('w:themeShade'), qn('w:themeTint')]:
                if at in el.attrib:
                    del el.attrib[at]
    print(f'  → {n} runs')


def s05_cover(doc):
    print('[05] Cover pages...')
    ch = 0
    for i, para in enumerate(doc.paragraphs[:30]):
        if i <= 14 or (15 <= i <= 24):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            for shd in pPr.findall(qn('w:shd')):
                pPr.remove(shd); ch += 1
    mc_ns = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
    for i, para in enumerate(doc.paragraphs[:15]):
        for mc in para._element.findall(f'{{{mc_ns}}}AlternateContent'):
            para._element.remove(mc); ch += 1
    bg = doc.element.find(qn('w:background'))
    if bg is not None:
        doc.element.remove(bg); ch += 1
    print(f'  → {ch} fixes')


def s06_page_numbers(doc):
    print('[06] Page numbering...')
    secs = doc.sections
    n = len(secs)
    last = n - 1

    for i, sec in enumerate(secs):
        sp = sec._sectPr
        for old in sp.findall(qn('w:pgNumType')):
            sp.remove(old)

        if i == 0:
            sec.different_first_page_header_footer = True
            for p in sec.first_page_footer.paragraphs:
                _clear_p(p._element)
            sec.footer.is_linked_to_previous = False
            fps = sec.footer.paragraphs
            fp = fps[0]._element if fps else OxmlElement('w:p')
            if not fps:
                sec.footer._element.append(fp)
            _clear_p(fp); _center(fp)
            for r in _page_field_runs():
                fp.append(r)
            pnt = OxmlElement('w:pgNumType')
            pnt.set(qn('w:fmt'), 'lowerRoman'); pnt.set(qn('w:start'), '0')
            sp.append(pnt)

        elif i < last:
            sec.footer.is_linked_to_previous = False
            fps = sec.footer.paragraphs
            fp = fps[0]._element if fps else OxmlElement('w:p')
            if not fps:
                sec.footer._element.append(fp)
            _clear_p(fp); _center(fp)
            for r in _page_field_runs():
                fp.append(r)
            pnt = OxmlElement('w:pgNumType')
            pnt.set(qn('w:fmt'), 'lowerRoman')
            sp.append(pnt)
            for extra in list(sec.footer.paragraphs[1:]):
                _del(extra._element)

        else:
            sec.footer.is_linked_to_previous = False
            fps = sec.footer.paragraphs
            fp = fps[0]._element if fps else OxmlElement('w:p')
            if not fps:
                sec.footer._element.append(fp)
            _clear_p(fp); _center(fp)
            for r in _page_field_runs():
                fp.append(r)
            pnt = OxmlElement('w:pgNumType')
            pnt.set(qn('w:fmt'), 'decimal'); pnt.set(qn('w:start'), '1')
            sp.append(pnt)
            for extra in list(sec.footer.paragraphs[1:]):
                _del(extra._element)

    print(f'  → {n} sections (cover + {last-1} Roman + 1 Arabic)')


def s07_tablo_nums(doc):
    print('[07] Tablo 0.x → 1.x...')
    ch = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if 'Tablo 0.' in run.text:
                run.text = run.text.replace('Tablo 0.', 'Tablo 1.'); ch += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if 'Tablo 0.' in run.text:
                            run.text = run.text.replace('Tablo 0.', 'Tablo 1.'); ch += 1
    print(f'  → {ch} refs')


def s08_kisaltmalar(doc):
    print('[08] KISALTMALAR...')
    idx = _find_idx(doc, 'KISALTMALAR')
    if idx is None:
        print('  ! not found'); return
    nxt = _next_h1(doc, idx)
    if nxt:
        _del_between(doc, idx, nxt)
    anchor = doc.paragraphs[idx]._element
    # insert blank line
    _insert_after(anchor, '')
    for abbr, meaning in reversed(ABBREVIATIONS):
        p = OxmlElement('w:p'); anchor.addnext(p)
        pPr = OxmlElement('w:pPr')
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'left'); tab.set(qn('w:pos'), '3600')
        tabs.append(tab); pPr.append(tabs); p.insert(0, pPr)
        _add_run(p, abbr, bold=True)
        rt = OxmlElement('w:r'); rt.append(OxmlElement('w:tab')); p.append(rt)
        _add_run(p, meaning)
    print(f'  → {len(ABBREVIATIONS)} entries')


def s09_toc(doc):
    print('[09] İÇİNDEKİLER (TOC field)...')
    idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'İÇİNDEKİLER' and 'Heading' in p.style.name:
            idx = i; break
    if idx is None:
        print('  ! not found'); return
    nxt = _next_h1(doc, idx)
    if nxt:
        _del_between(doc, idx, nxt)
    anchor = doc.paragraphs[idx]._element
    # "Sayfa" right
    ps = _insert_after(anchor)
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'right'); pPr.append(jc)
    ps.insert(0, pPr)
    _add_run(ps, 'Sayfa', bold=True)
    # TOC field
    pt = _insert_after(ps)
    for ft, content in [('begin',None),('instr',' TOC \\o "1-3" \\h \\z \\u '),
                         ('sep',None),('disp','İçindekiler tablosunu güncellemek için sağ tıklayıp "Alanı Güncelle" seçin.'),
                         ('end',None)]:
        r = OxmlElement('w:r'); _set_run_fmt(r)
        if ft == 'begin':
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin'); r.append(fc)
        elif ft == 'instr':
            it = OxmlElement('w:instrText')
            it.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            it.text = content; r.append(it)
        elif ft == 'sep':
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'separate'); r.append(fc)
        elif ft == 'disp':
            t = OxmlElement('w:t'); t.text = content; r.append(t)
        elif ft == 'end':
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'end'); r.append(fc)
        pt.append(r)
    print('  → TOC field (update in Word: Ctrl+A → F9)')


def _rebuild_list(doc, heading, prefix):
    idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == heading and 'Heading' in p.style.name:
            idx = i; break
    if idx is None:
        return 0
    nxt = _next_h1(doc, idx)
    if nxt:
        _del_between(doc, idx, nxt)
    # find body captions
    caps = []; seen = set(); past_lists = False
    for p in doc.paragraphs:
        if p.text.strip() == 'GİRİŞ' and 'Heading' in p.style.name:
            past_lists = True
        if not past_lists:
            continue
        txt = p.text.strip()
        if re.match(rf'^{re.escape(prefix)}\s+\d+\.\d+', txt):
            key = re.sub(r'\s+', ' ', txt).rstrip('.')
            if key not in seen:
                seen.add(key); caps.append(key)
    anchor = doc.paragraphs[idx]._element
    ps = _insert_after(anchor)
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'right'); pPr.append(jc)
    ps.insert(0, pPr)
    _add_run(ps, 'Sayfa', bold=True)
    prev = ps
    for cap in caps:
        np = _insert_after(prev, cap)
        pPr = np.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr'); np.insert(0, pPr)
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right'); tab.set(qn('w:leader'), 'dot')
        tab.set(qn('w:pos'), '9072')
        tabs.append(tab); pPr.append(tabs)
        prev = np
    return len(caps)


def s10_tables_list(doc):
    print('[10] TABLOLAR LİSTESİ...')
    n = _rebuild_list(doc, 'TABLOLAR LİSTESİ', 'Tablo')
    print(f'  → {n} entries')


def s11_figures_list(doc):
    print('[11] ŞEKİLLER LİSTESİ...')
    n = _rebuild_list(doc, 'ŞEKİLLER LİSTESİ', 'Şekil')
    print(f'  → {n} entries')


def s12_forbidden(doc):
    print('[12] Forbidden words...')
    ch = 0
    for para in doc.paragraphs:
        for w in FORBIDDEN:
            if w in para.text:
                for run in para.runs:
                    if w in run.text:
                        run.text = run.text.replace(w, ''); ch += 1
    print(f'  → {ch} removed' if ch else '  → clean')


def s13_numbers(doc):
    print('[13] Numerical consistency...')
    ch = 0
    for para in doc.paragraphs:
        txt = para.text
        if '210 test' in txt or '210 test' in txt.lower():
            for run in para.runs:
                if '210' in run.text:
                    run.text = run.text.replace('210', '219'); ch += 1
        if ("203" in txt) and ('geçmiş' in txt or 'passed' in txt.lower()):
            for run in para.runs:
                if '203' in run.text:
                    run.text = run.text.replace('203', '212'); ch += 1
    print(f'  → {ch} fixes')


def s14_noproof(doc):
    print('[14] noProof...')
    body = doc.element.body; n = 0
    for r_el in body.iter(qn('w:r')):
        t_el = r_el.find(qn('w:t'))
        if t_el is None or not t_el.text:
            continue
        if any(term in t_el.text for term in TECH_TERMS):
            rPr = r_el.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr'); r_el.insert(0, rPr)
            if rPr.find(qn('w:noProof')) is None:
                rPr.append(OxmlElement('w:noProof')); n += 1
    print(f'  → {n} runs')


def s15_language(doc):
    print('[15] Language...')
    body = doc.element.body
    for r_el in body.iter(qn('w:r')):
        rPr = r_el.find(qn('w:rPr'))
        if rPr is None:
            continue
        lang = rPr.find(qn('w:lang'))
        if lang is None:
            lang = OxmlElement('w:lang'); rPr.append(lang)
        lang.set(qn('w:val'), 'tr-TR')
    in_abs = False
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt == 'ABSTRACT':
            in_abs = True; continue
        if in_abs and para.style.name.startswith('Heading 1') and txt:
            break
        if in_abs:
            for run in para.runs:
                rPr = run._element.find(qn('w:rPr'))
                if rPr is not None:
                    lang = rPr.find(qn('w:lang'))
                    if lang is not None:
                        lang.set(qn('w:val'), 'en-US')
    print('  → tr-TR / ABSTRACT en-US')


def s16_hyperlinks(doc):
    print('[16] Hyperlinks...')
    body = doc.element.body; n = 0
    for hl in body.iter(qn('w:hyperlink')):
        for r_el in hl.iter(qn('w:r')):
            _set_run_fmt(r_el)
            rPr = r_el.find(qn('w:rPr'))
            if rPr is not None:
                for rs in rPr.findall(qn('w:rStyle')):
                    rPr.remove(rs)
            n += 1
    print(f'  → {n} runs')


def s17_captions(doc):
    print('[17] Caption formatting...')
    n = 0
    for para in doc.paragraphs:
        txt = para.text.strip()
        if re.match(r'^(Tablo|Şekil)\s+\d+\.\d+', txt):
            for run in para.runs:
                run.italic = False
                run.font.size = Pt(12)
                run.font.name = FONT
                run.font.color.rgb = RGBColor(0, 0, 0)
                n += 1
    print(f'  → {n} runs')


def s18_quality(doc):
    print('\n[QC] Quality Check')
    ok = True
    # Fonts
    bad = set()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name and run.font.name != FONT:
                bad.add(run.font.name)
    if bad:
        print(f'  ⚠ Fonts: {bad}'); ok = False
    else:
        print('  ✓ All fonts TNR')
    # Sizes
    bad = set()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size and run.font.size != Pt(12):
                bad.add(str(run.font.size))
    if bad:
        print(f'  ⚠ Sizes: {bad}'); ok = False
    else:
        print('  ✓ All 12pt')
    # Colors
    bad = set()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.color and run.font.color.rgb and run.font.color.rgb != RGBColor(0,0,0):
                bad.add(str(run.font.color.rgb))
    if bad:
        print(f'  ⚠ Colors: {bad}'); ok = False
    else:
        print('  ✓ All black')
    # Tablo 0.x
    if any('Tablo 0.' in p.text for p in doc.paragraphs):
        print('  ⚠ Tablo 0.x remains'); ok = False
    else:
        print('  ✓ No Tablo 0.x')
    # Forbidden
    fw = []
    for p in doc.paragraphs:
        for w in FORBIDDEN:
            if w in p.text:
                fw.append(w)
    if fw:
        print(f'  ⚠ Forbidden: {set(fw)}'); ok = False
    else:
        print('  ✓ No forbidden words')
    if ok:
        print('\n  ✓ ALL CHECKS PASSED')
    return ok


# ── Main ───────────────────────────────────────────────────────

def main():
    print('=' * 60)
    print('DYTOPIA THESIS DOCX FORMATTER')
    print('=' * 60)

    s01_backup()
    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(str(OUTPUT))

    s02_doc_defaults(doc)
    s03_styles(doc)
    s04_runs(doc)
    s05_cover(doc)
    s06_page_numbers(doc)
    s07_tablo_nums(doc)
    s08_kisaltmalar(doc)
    s09_toc(doc)
    s10_tables_list(doc)
    s11_figures_list(doc)
    s12_forbidden(doc)
    s13_numbers(doc)
    s14_noproof(doc)
    s15_language(doc)
    s16_hyperlinks(doc)
    s17_captions(doc)

    print('\n[SAVE]...')
    doc.save(str(OUTPUT))
    print(f'  → {OUTPUT.name}')

    print('\n[VERIFY] Re-opening...')
    doc2 = Document(str(OUTPUT))
    s18_quality(doc2)

    print('\n' + '=' * 60)
    print('TAMAMLANDI')
    print('=' * 60)
    print('\nManuel kontrol gereken maddeler:')
    print('  1. Word\'de açıp Ctrl+A → F9 ile tüm alanları güncelleyin')
    print('  2. İçindekiler sayfa numaralarını kontrol edin')
    print('  3. Tablo başlıkları tablonun üstünde mi doğrulayın')
    print('  4. Şekil başlıkları şeklin altında mı doğrulayın')
    print('  5. Görsellerin sayfa sınırları içinde kaldığını kontrol edin')


if __name__ == '__main__':
    main()
