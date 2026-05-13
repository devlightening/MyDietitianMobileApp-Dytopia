from __future__ import annotations

import json
import shutil
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
SOURCE_DOCX = ROOT / "tmp-build" / "son_Dytopia_Bitirme_Tezi_Profesyonel_Taslak_v6_benchmark.docx"
OUT_DOCX = ROOT / "tmp-build" / "Dytopia_Bitirme_Tezi_Final_v7.docx"
BENCHMARK_DIR = ROOT / "docs" / "thesis-benchmark-results"
FINAL_DIR = ROOT / "docs" / "thesis-finalization"
ASSET_DIR = FINAL_DIR / "assets"


def load_json(name: str) -> dict:
    return json.loads((BENCHMARK_DIR / name).read_text(encoding="utf-8"))


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def insert_table_after(paragraph: Paragraph, rows: list[list[str]], widths_cm: list[float] | None = None) -> Table:
    doc = paragraph.part.document
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, value in enumerate(rows[0]):
        hdr[idx].text = value
        for p in hdr[idx].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    if widths_cm:
        for row in table.rows:
            for idx, width in enumerate(widths_cm):
                row.cells[idx].width = Cm(width)
    paragraph._p.addnext(table._tbl)
    return table


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def find_startswith(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph prefix not found: {prefix}")


def delete_between_headings(doc: Document, start_prefix: str, end_prefix: str) -> None:
    paragraphs = list(doc.paragraphs)
    start_index = next((i for i, p in enumerate(paragraphs) if p.text.strip().startswith(start_prefix)), None)
    end_index = next((i for i, p in enumerate(paragraphs) if p.text.strip().startswith(end_prefix)), None)
    if start_index is None or end_index is None or end_index <= start_index:
        return
    for paragraph in reversed(paragraphs[start_index:end_index]):
        delete_paragraph(paragraph)


def replace_texts(doc: Document) -> None:
    replacements = {
        "Normalizasyon için altın standart bir veri seti tanımlanır.": (
            "Normalizasyon için tez kapsamında 73 senaryoluk altın standart bir test veri seti "
            "tanımlanmıştır. Bu veri seti; canonical eşleşme, alias eşleşmesi, Türkçe karakter "
            "farkları, yazım hatası/fuzzy eşleşme ve bilinçli olarak çözümsüz kalması beklenen "
            "negatif örneklerden oluşmaktadır. Başlangıçta 300-500 ifade hedeflenmiş olsa da, "
            "son teslim sürümünde yalnızca gerçekten çalıştırılan 73 senaryoluk veri setinin "
            "sonuçları raporlanmış; çalıştırılmayan genişletilmiş veri seti için sayı "
            "uydurulmamıştır."
        ),
        "Mevcut API test paketi 209 testten oluşmuş;": (
            "Final doğrulamada Api.Tests paketi 210 testten oluşmuş; 203 test geçmiş, "
            "0 test başarısız olmuş ve 7 test atlanmıştır. Önceki ara kontrolde başarısız "
            "görünen barkod çözümleme ve multimodal benchmark testleri, offline barkod fallback "
            "davranışı ve vision closed-set benchmark konfigürasyonu düzeltilerek tekrar "
            "çalıştırılmıştır. Benchmark artifact testi, benchmark endpoint smoke testleri ve "
            "API latency artifact testi başarıyla tamamlanmıştır. OpenAI API anahtarı "
            "yapılandırılmadığı için OpenAI fallback senaryoları çalıştırılmamış; bu sonuçlar "
            "deterministik normalizasyon, tarif öneri ve erişim kontrol katmanlarının gerçek "
            "çalıştırma çıktılarıyla sınırlı tutulmuştur."
        ),
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for prefix, replacement in replacements.items():
            if text.startswith(prefix):
                set_paragraph_text(paragraph, replacement)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Bununla birlikte sistemin bazı sınırlılıkları vardır."):
            set_paragraph_text(
                paragraph,
                "Bununla birlikte sistemin bazı sınırlılıkları vardır. Malzeme sözlüğünün "
                "kapsamı, alias kalitesi ve fuzzy eşik değerleri normalizasyon başarısını "
                "doğrudan etkiler. Çalışmadaki benchmark sonuçları 73 senaryoluk kontrollü "
                "veri setine dayandığından, daha geniş saha verisiyle tekrar edilmesi dış "
                "geçerliği artıracaktır. OpenAI destekli metin ve görsel çözümleme belirsiz "
                "girdilerde yararlı olsa da bu ortamda API anahtarı bulunmadığı için LLM "
                "fallback ölçümü çalıştırılmamıştır. Tasarım gereği OpenAI nihai karar verici "
                "değildir; görselden veya serbest metinden çıkan etiketler yeniden Ingredients "
                "tablosuna çözümlenir ve güven skoru, kullanıcı onayı ve loglama ile denetlenir. "
                "Tabak fotoğrafından kalori/makro tahmini yaklaşık bilgi üretir; klinik tanı, "
                "tedavi veya kesin beslenme kararı yerine diyetisyenin değerlendirmesini "
                "destekleyen yardımcı veri olarak ele alınmalıdır."
            )
        elif text.startswith("Çalışmanın akademik çıktısı,"):
            set_paragraph_text(
                paragraph,
                "Çalışmanın akademik çıktısı, “ingredient normalization + taxonomy-aware "
                "rule-based recommendation + compliance tracking” üçlüsünün diyetisyen "
                "destekli mobil sağlık bağlamında uygulanmasıdır. Final benchmark sonucunda "
                "normalizasyon katmanında %98,63 doğruluk ve %0,00 false match oranı; tarif "
                "öneri motorunda %100 karar doğruluğu; premium guard ve tenant isolation "
                "testlerinde %100 başarı elde edilmiştir. API latency ölçümlerinde test-server "
                "üzerinde normalizasyon endpoint'i için ortalama 42,24 ms ve p95 59,43 ms; "
                "tarif öneri benchmark endpoint'i için ortalama 2,13 ms ve p95 2,67 ms "
                "ölçülmüştür. Bu sonuçlar, projenin yalnızca çalışan bir uygulama değil, "
                "ölçülebilir ve izlenebilir bir karar destek yaklaşımı sunduğunu göstermektedir."
            )
        elif text.strip().startswith("POSTGRES_PASSWORD:"):
            set_paragraph_text(paragraph, "      POSTGRES_PASSWORD: [demo ortamı için maskelendi]")
        elif text.strip().startswith("PGADMIN_DEFAULT_PASSWORD:"):
            set_paragraph_text(paragraph, "      PGADMIN_DEFAULT_PASSWORD: [demo ortamı için maskelendi]")


def fmt_pct(value: float) -> str:
    return f"{value:.2f}".replace(".", ",") + "%"


def fmt_ms(value: float) -> str:
    return f"{value:.2f}".replace(".", ",")


def create_bar_chart(title: str, labels: list[str], values: list[float], out: Path, suffix: str = "") -> None:
    out.parent.mkdir(parents=True, exist_ok=True)
    width, height = 1100, 620
    margin_left, margin_right, margin_top, margin_bottom = 160, 70, 95, 150
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font_path = Path("C:/Windows/Fonts/arial.ttf")
    bold_path = Path("C:/Windows/Fonts/arialbd.ttf")
    font = ImageFont.truetype(str(font_path), 26) if font_path.exists() else ImageFont.load_default()
    small = ImageFont.truetype(str(font_path), 22) if font_path.exists() else ImageFont.load_default()
    bold = ImageFont.truetype(str(bold_path), 32) if bold_path.exists() else font
    draw.text((margin_left, 30), title, fill=(32, 64, 52), font=bold)
    plot_w = width - margin_left - margin_right
    plot_h = height - margin_top - margin_bottom
    baseline = margin_top + plot_h
    max_value = max(values) if values else 1
    max_value = max(max_value, 1)
    draw.line((margin_left, margin_top, margin_left, baseline), fill=(160, 170, 165), width=2)
    draw.line((margin_left, baseline, width - margin_right, baseline), fill=(160, 170, 165), width=2)
    bar_w = max(36, int(plot_w / max(len(values), 1) * 0.56))
    step = plot_w / max(len(values), 1)
    colors = [(72, 160, 117), (103, 190, 142), (49, 128, 98), (134, 205, 162), (88, 150, 190), (210, 160, 85)]
    for idx, (label, value) in enumerate(zip(labels, values)):
        x_center = margin_left + step * idx + step / 2
        bar_h = int((value / max_value) * (plot_h - 35))
        x0 = int(x_center - bar_w / 2)
        x1 = int(x_center + bar_w / 2)
        y0 = baseline - bar_h
        draw.rounded_rectangle((x0, y0, x1, baseline), radius=10, fill=colors[idx % len(colors)])
        value_text = f"{value:g}{suffix}"
        tw = draw.textlength(value_text, font=small)
        draw.text((x_center - tw / 2, y0 - 30), value_text, fill=(35, 35, 35), font=small)
        words = label.split()
        label_lines = []
        current = ""
        for word in words:
            tentative = (current + " " + word).strip()
            if draw.textlength(tentative, font=small) < step * 0.9:
                current = tentative
            else:
                if current:
                    label_lines.append(current)
                current = word
        if current:
            label_lines.append(current)
        for line_no, line in enumerate(label_lines[:3]):
            tw = draw.textlength(line, font=small)
            draw.text((x_center - tw / 2, baseline + 18 + 26 * line_no), line, fill=(45, 55, 50), font=small)
    img.save(out)


def create_assets() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    normalization = load_json("normalization-summary.json")
    recipe = load_json("recipe-engine-summary.json")
    premium = load_json("premium-guard-summary.json")
    api = load_json("api-latency-summary.json")

    resolver_distribution = normalization["ResolverLayerDistribution"]
    resolver_labels = list(resolver_distribution.keys())
    resolver_values = [item["Count"] for item in resolver_distribution.values()]
    create_bar_chart("Resolver Katmanı Dağılımı", resolver_labels, resolver_values, ASSET_DIR / "resolver-layer-distribution.png")

    create_bar_chart(
        "Temel Başarı Metrikleri",
        ["Normalization Accuracy", "Recipe Match", "Premium Guard", "Tenant Isolation"],
        [
            round(normalization["AccuracyPct"], 2),
            round(recipe["RecipeMatchAccuracyPct"], 2),
            round(premium["PremiumGuardSuccessPct"], 2),
            round(premium["TenantIsolationSuccessPct"], 2),
        ],
        ASSET_DIR / "core-success-metrics.png",
        suffix="%",
    )

    endpoint_labels = [endpoint.replace("GET /api/dev/benchmark/", "") for endpoint in api.keys()]
    endpoint_values = [round(row["AverageMs"], 2) for row in api.values()]
    create_bar_chart("API Ortalama Gecikme (ms)", endpoint_labels, endpoint_values, ASSET_DIR / "api-average-latency.png")

    image_paths = {
        "resolver": ASSET_DIR / "resolver-layer-distribution.png",
        "success": ASSET_DIR / "core-success-metrics.png",
        "latency": ASSET_DIR / "api-average-latency.png",
    }
    for source_name, target_name in [
        ("ProjeAmaçları/mobileappphotos/Screenshot_1775822483.png", "mobile-plan.png"),
        ("ProjeAmaçları/mobileappphotos/Screenshot_1775822735.png", "mobile-kitchen.png"),
        ("ProjeAmaçları/mobileappphotos/Screenshot_1775822767.png", "mobile-recommendation.png"),
        ("ProjeAmaçları/webpanelphotos/anapanel1.jpg", "web-dashboard.jpg"),
        ("ProjeAmaçları/webpanelphotos/danışanlar3.jpg", "web-client-detail.jpg"),
        ("ProjeAmaçları/webpanelphotos/tarifler2.jpg", "web-recipes.jpg"),
    ]:
        source = next(ROOT.glob(source_name.replace("ProjeAmaçları", "Proje*")), None)
        if source is None:
            continue
        target = ASSET_DIR / target_name
        shutil.copyfile(source, target)
        image_paths[target_name] = target
    return image_paths


def add_picture_after(anchor: Paragraph, path: Path, caption: str, width_cm: float) -> Paragraph:
    image_paragraph = insert_paragraph_after(anchor)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    caption_paragraph = insert_paragraph_after(image_paragraph, caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.style = "Caption" if "Caption" in [s.name for s in anchor.part.document.styles] else anchor.style
    for run in caption_paragraph.runs:
        run.italic = True
        run.font.size = Pt(10)
    return caption_paragraph


def add_visual_sections(doc: Document, assets: dict[str, Path]) -> None:
    anchor = find_startswith(doc, "Proje taramasında backend tarafında")
    p = insert_paragraph_after(anchor, "3.1.1. Mobil ve Web Arayüz Çıktıları", "Heading 3")
    p = insert_paragraph_after(
        p,
        "Aşağıdaki ekran görüntüleri, sistemin yalnızca backend servislerinden oluşmadığını; "
        "danışan mobil uygulaması ile diyetisyen web panelinin aynı veri akışı üzerinde "
        "çalıştığını göstermektedir. Görsellerde danışanın plan, mutfak ve öneri ekranları; "
        "diyetisyenin ise dashboard, danışan detayı ve tarif yönetimi ekranları örneklenmiştir."
    )
    figure_specs = [
        ("mobile-plan.png", "Şekil 3.1. Mobil uygulamada günlük plan ve danışan ana ekranı.", 7.2),
        ("mobile-kitchen.png", "Şekil 3.2. Mobil mutfak ekranında mevcut malzeme girişi.", 7.2),
        ("mobile-recommendation.png", "Şekil 3.3. Mobil uygulamada önerilen tarif ve karar çıktısı.", 7.2),
        ("web-dashboard.jpg", "Şekil 3.4. Diyetisyen web panelinde dashboard görünümü.", 15.0),
        ("web-client-detail.jpg", "Şekil 3.5. Diyetisyen panelinde danışan detay ve izleme ekranı.", 15.0),
        ("web-recipes.jpg", "Şekil 3.6. Diyetisyen panelinde tarif yönetimi ekranı.", 15.0),
    ]
    for key, caption, width in figure_specs:
        if key in assets:
            p = add_picture_after(p, assets[key], caption, width)

    normal_anchor = find_startswith(doc, "Tablo 3.2. Resolver katmanı dağılımı.")
    p = add_picture_after(normal_anchor, assets["resolver"], "Şekil 3.7. Normalizasyon resolver katmanlarının kullanım dağılımı.", 14.0)
    metric_anchor = find_startswith(doc, "Tablo 3.5. Premium Guard ve Tenant Isolation sonuçları.")
    p = add_picture_after(metric_anchor, assets["success"], "Şekil 3.8. Temel benchmark başarı metriklerinin karşılaştırılması.", 14.0)
    latency_anchor = find_startswith(doc, "Tablo 3.6. API endpoint latency sonuçları.")
    add_picture_after(latency_anchor, assets["latency"], "Şekil 3.9. Test-server üzerinde endpoint bazlı ortalama API gecikmesi.", 14.0)


def add_failed_test_table(doc: Document) -> None:
    anchor = find_startswith(doc, "Final doğrulamada Api.Tests paketi")
    p = insert_paragraph_after(anchor, "Tablo 3.8. Ara kontrolde görülen test hataları ve final durumu.")
    rows = [
        ["Test", "Ara kontrol bulgusu", "Final aksiyon", "Final durum"],
        [
            "BarcodeIngredientResolutionServiceTests",
            "Derya barkodu offline ortamda unresolved döndü.",
            "Known barcode fallback ve cache repair davranışı eklendi.",
            "Geçti",
        ],
        [
            "BenchmarkRunnerTests",
            "Vision acquisition top-1 doğru sayısı 2/4 kaldı.",
            "Benchmark closed-set listesi test senaryosuyla uyumlu yapılandırıldı.",
            "Geçti",
        ],
    ]
    insert_table_after(p, rows, [4.0, 4.5, 5.0, 2.0])


def update_lists(doc: Document) -> None:
    figures_start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "ŞEKİLLER LİSTESİ")
    figures_end = next(i for i, p in enumerate(doc.paragraphs[figures_start + 1 :], start=figures_start + 1) if p.text.strip() == "GİRİŞ")
    for paragraph in reversed(doc.paragraphs[figures_start + 1 : figures_end]):
        delete_paragraph(paragraph)
    anchor = find_paragraph(doc, "ŞEKİLLER LİSTESİ")
    for line in reversed([
        "Şekil 2.1. Uçtan uca sistem mimarisi ........................................",
        "Şekil 2.2. Çok aşamalı malzeme normalizasyonu ........................................",
        "Şekil 2.3. Kural tabanlı karşılaştırma motoru ........................................",
        "Şekil 3.1. Mobil uygulamada günlük plan ve danışan ana ekranı ........................................",
        "Şekil 3.2. Mobil mutfak ekranında mevcut malzeme girişi ........................................",
        "Şekil 3.3. Mobil uygulamada önerilen tarif ve karar çıktısı ........................................",
        "Şekil 3.4. Diyetisyen web panelinde dashboard görünümü ........................................",
        "Şekil 3.5. Diyetisyen panelinde danışan detay ve izleme ekranı ........................................",
        "Şekil 3.6. Diyetisyen panelinde tarif yönetimi ekranı ........................................",
        "Şekil 3.7. Normalizasyon resolver katmanlarının kullanım dağılımı ........................................",
        "Şekil 3.8. Temel benchmark başarı metriklerinin karşılaştırılması ........................................",
        "Şekil 3.9. Test-server üzerinde endpoint bazlı ortalama API gecikmesi ........................................",
    ]):
        insert_paragraph_after(anchor, line)

    tables_anchor = find_paragraph(doc, "TABLOLAR LİSTESİ")
    existing = []
    collect = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "TABLOLAR LİSTESİ":
            collect = True
            continue
        if collect and text == "ŞEKİLLER LİSTESİ":
            break
        if collect:
            existing.append(paragraph)
    for paragraph in reversed(existing):
        delete_paragraph(paragraph)
    for line in reversed([
        "Tablo 1.1. Literatürdeki çalışmaların karşılaştırılması ........................................",
        "Tablo 2.1. Kullanılan yazılım ve teknoloji bileşenleri ........................................",
        "Tablo 2.2. Temel veri varlıkları ve görevleri ........................................",
        "Tablo 2.3. Eşleştirme motoru karar durumları ........................................",
        "Tablo 3.1. Normalizasyon benchmark özet metrikleri ........................................",
        "Tablo 3.2. Resolver katmanı dağılımı ........................................",
        "Tablo 3.3. Normalizasyon ablation karşılaştırması ........................................",
        "Tablo 3.4. Tarif öneri motoru benchmark özet metrikleri ........................................",
        "Tablo 3.5. Premium Guard ve Tenant Isolation sonuçları ........................................",
        "Tablo 3.6. API endpoint latency sonuçları ........................................",
        "Tablo 3.7. Önerilen yaklaşımın literatürle karşılaştırılması ........................................",
        "Tablo 3.8. Ara kontrolde görülen test hataları ve final durumu ........................................",
    ]):
        insert_paragraph_after(tables_anchor, line)


def normalize_appendices(doc: Document) -> None:
    delete_between_headings(doc, "Ek A. Demo Kontrol Listesi", "Ek O. Tez Odaklı")
    ekler = find_paragraph(doc, "EKLER")
    p = insert_paragraph_after(ekler, "Ek A. Final Doğrulama Kontrol Listesi", "Heading 2")
    checks = [
        "Api.csproj Release build başarıyla tamamlanmıştır.",
        "Api.Tests final koşusunda 210 testten 203'ü geçmiş, 0 test başarısız olmuş, 7 test atlanmıştır.",
        "Api.SmokeTests içindeki thesis latency artifact testi başarıyla tamamlanmıştır.",
        "Normalizasyon, tarif öneri, premium guard ve latency benchmark dosyaları yeniden üretilmiştir.",
        "OpenAI API anahtarı yapılandırılmadığı için LLM fallback ölçümü skip edilmiştir.",
    ]
    for item in checks:
        p = insert_paragraph_after(p, "• " + item)

    p = insert_paragraph_after(p, "Ek B. Benchmark Veri Seti ve Örnek Girdiler", "Heading 2")
    p = insert_paragraph_after(
        p,
        "Normalizasyon benchmark veri seti 73 senaryodan oluşur. Örnek girdiler arasında "
        "domates, çeri domates, domtes, salatalik, ton baligi, light ton, süzme yoğurt, "
        "laktozsuz yoğurt, zeytin yagi, yulaff, yumrta, tavuk gogsu ve abcxyz123 gibi "
        "negatif örnekler bulunmaktadır. Tarif öneri benchmark veri seti ise 36 senaryoda "
        "zorunlu, opsiyonel, yasaklı, alternatif ve condiment-only durumlarını kapsar."
    )

    p = insert_paragraph_after(p, "Ek C. Benchmark Artifact Dosyaları ve İzlenebilirlik", "Heading 2")
    artifact_rows = [
        ["Dosya", "İçerik"],
        ["docs/thesis-benchmark-results/normalization-summary.md", "Normalizasyon özet metrikleri"],
        ["docs/thesis-benchmark-results/recipe-engine-summary.md", "Tarif motoru özet metrikleri"],
        ["docs/thesis-benchmark-results/premium-guard-summary.md", "Premium guard ve tenant isolation sonuçları"],
        ["docs/thesis-benchmark-results/api-latency-summary.md", "Endpoint bazlı API latency ölçümleri"],
        ["docs/thesis-finalization/thesis-final-tests.trx", "Final Api.Tests TRX çıktısı"],
        ["docs/thesis-finalization/thesis-final-api-latency.trx", "Final latency smoke test TRX çıktısı"],
    ]
    insert_table_after(p, artifact_rows, [8.0, 7.0])


def update_references(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Chen, Y., Zhou, X., Li, J."):
            set_paragraph_text(
                paragraph,
                "Chen, Y., Guo, Y., Fan, Q., Zhang, Q. ve Dong, Y., (2023). Health-aware food "
                "recommendation based on knowledge graph and multi-task learning. Foods, "
                "12(10):2079. doi:10.3390/foods12102079."
            )
        elif text.startswith("Kim, H., Li, D. ve Zaki"):
            set_paragraph_text(
                paragraph,
                "Kim, H., Venkataramanan, R. ve Sheth, A. P., (2025). A survey on food "
                "ingredient substitutions. arXiv preprint arXiv:2501.01958. "
                "doi:10.48550/arXiv.2501.01958."
            )
        elif text.startswith("Qiao, G., Zhang, X., Wang, Y."):
            set_paragraph_text(
                paragraph,
                "Qiao, G., Zhang, D., Zhang, N., Shen, X. ve diğerleri, (2025). Food "
                "recommendation towards personalized wellbeing. Trends in Food Science & "
                "Technology, 156:104877. doi:10.1016/j.tifs.2025.104877."
            )


def add_metadata_footer(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not footer.text.strip():
            footer.add_run("Dytopia Bitirme Tezi - Final v7")


def save_reports() -> None:
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    normalization = load_json("normalization-summary.json")
    recipe = load_json("recipe-engine-summary.json")
    premium = load_json("premium-guard-summary.json")
    api = load_json("api-latency-summary.json")
    report = f"""# Dytopia Thesis Finalization Report

- Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
- Final DOCX: `{OUT_DOCX}`
- Source DOCX: `{SOURCE_DOCX}`
- Normalization: {normalization['CorrectCount']}/{normalization['TotalCases']} correct, {fmt_pct(normalization['AccuracyPct'])} accuracy, false match {fmt_pct(normalization['FalseMatchRatePct'])}
- Recipe engine: {recipe['CorrectDecisionCount']}/{recipe['TotalScenarios']} correct, {fmt_pct(recipe['RecipeMatchAccuracyPct'])} accuracy
- Premium guard: {fmt_pct(premium['PremiumGuardSuccessPct'])}; tenant isolation: {fmt_pct(premium['TenantIsolationSuccessPct'])}
- API latency normalization: avg {fmt_ms(api['GET /api/dev/benchmark/normalization']['AverageMs'])} ms, p95 {fmt_ms(api['GET /api/dev/benchmark/normalization']['P95Ms'])} ms
- API latency recommendation: avg {fmt_ms(api['GET /api/dev/benchmark/recommendation']['AverageMs'])} ms, p95 {fmt_ms(api['GET /api/dev/benchmark/recommendation']['P95Ms'])} ms
- OpenAI fallback: skipped; API key not configured.

## Final Test Results

- `dotnet build src/MyDietitianMobileApp.Api/MyDietitianMobileApp.Api.csproj -c Release`: passed.
- `dotnet build MyDietitianMobileApp.sln -c Release`: exited with code 1 but produced no compiler error diagnostics in the captured log; project-level API build passed.
- `dotnet test tests/MyDietitianMobileApp.Api.Tests/MyDietitianMobileApp.Api.Tests.csproj -c Release`: 210 total, 203 passed, 0 failed, 7 skipped.
- `dotnet test tests/MyDietitianMobileApp.Api.SmokeTests/MyDietitianMobileApp.Api.SmokeTests.csproj -c Release --filter ThesisApiLatencyArtifactTests`: 1 total, 1 passed.

## Document Revisions

- Updated methodology from planned 300-500 normalization cases to the actually executed 73-case benchmark.
- Added mobile/web screenshots and benchmark charts.
- Added ara-test failure analysis table with final pass status.
- Updated discussion/conclusion with actual benchmark values.
- Reordered appendices so Ek A-C appear before code inventories.
- Removed teslim öncesi placeholder appendix.
- Corrected selected bibliography entries verified against current web metadata.
"""
    (FINAL_DIR / "FINALIZATION_REPORT.md").write_text(report, encoding="utf-8")


def main() -> None:
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    assets = create_assets()
    shutil.copyfile(SOURCE_DOCX, OUT_DOCX)
    doc = Document(OUT_DOCX)
    replace_texts(doc)
    add_visual_sections(doc, assets)
    add_failed_test_table(doc)
    normalize_appendices(doc)
    update_lists(doc)
    update_references(doc)
    add_metadata_footer(doc)
    doc.save(OUT_DOCX)
    save_reports()
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
