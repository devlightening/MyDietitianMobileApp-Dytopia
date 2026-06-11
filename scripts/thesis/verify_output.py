import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
import docx
from docx.shared import Pt, RGBColor
from lxml import etree

doc = docx.Document(r'c:\Users\hy971\source\repos\MyDietitianMobileApp\Dytopia_Bitirme_Tezi_TarsusUni_Final_Duzenlenmis.docx')

print('=== FONTS CHECK ===')
fonts = set()
for p in doc.paragraphs:
    for r in p.runs:
        if r.font.name: fonts.add(r.font.name)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if r.font.name: fonts.add(r.font.name)
print(f'Fonts: {fonts}')

print('\n=== SIZE CHECK ===')
sizes = set()
for p in doc.paragraphs:
    for r in p.runs:
        if r.font.size: sizes.add(r.font.size)
print(f'Sizes (EMU): {sizes}')

print('\n=== COLOR CHECK ===')
colors = set()
for p in doc.paragraphs:
    for r in p.runs:
        if r.font.color and r.font.color.rgb:
            colors.add(str(r.font.color.rgb))
print(f'Colors: {colors}')

print('\n=== SECTIONS & FOOTERS ===')
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
for i, sec in enumerate(doc.sections):
    sp = sec._sectPr
    pnt = sp.find(f'{{{W}}}pgNumType')
    fmt = pnt.get(f'{{{W}}}fmt') if pnt is not None else 'N/A'
    start = pnt.get(f'{{{W}}}start', 'cont') if pnt is not None else 'N/A'
    f_xml = etree.tostring(sec.footer.paragraphs[0]._element).decode() if sec.footer.paragraphs else ''
    has_field = 'fldChar' in f_xml
    print(f'S{i}: fmt={fmt} start={start} diff_first={sec.different_first_page_header_footer} has_page_field={has_field}')

print('\n=== TABLO 0.x CHECK ===')
found = False
for p in doc.paragraphs:
    if 'Tablo 0.' in p.text:
        print(f'  PROBLEM: {p.text[:100]}')
        found = True
if not found:
    print('  CLEAN - no Tablo 0.x')

print('\n=== KISALTMALAR (first 5) ===')
k_idx = None
count = 0
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'KISALTMALAR':
        k_idx = i
    elif k_idx and i > k_idx and p.text.strip() and count < 5:
        if p.style.name.startswith('Heading'):
            break
        print(f'  {p.text.strip()[:70]}')
        count += 1

print('\n=== KEY STYLE DEFINITIONS ===')
for name in ['Normal', 'Heading 1', 'Heading 2', 'Caption', 'Caption TR']:
    try:
        st = doc.styles[name]
        f = st.font
        c = f.color.rgb if f.color and f.color.rgb else None
        print(f'  {name}: font={f.name} size={f.size} color={c}')
    except:
        pass

print('\n=== FORBIDDEN WORDS ===')
forbidden = ['DTHybrit', 'TarsusFer', 'Saim Gezer', 'MovieLens', 'SASRec', 'GRU4Rec', 'DeepFM']
found_any = False
for p in doc.paragraphs:
    for w in forbidden:
        if w in p.text:
            print(f'  FOUND: {w}')
            found_any = True
if not found_any:
    print('  CLEAN')

print('\n=== TABLE/FIGURE LISTS REBUILT ===')
tl_idx = None
sl_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'TABLOLAR LISTESI' or p.text.strip() == 'TABLOLAR L\u0130STES\u0130':
        tl_idx = i
    if p.text.strip() == 'SEKILLER LISTESI' or p.text.strip() == '\u015eEK\u0130LLER L\u0130STES\u0130':
        sl_idx = i
if tl_idx:
    count = 0
    for i in range(tl_idx+1, min(tl_idx+20, len(doc.paragraphs))):
        if doc.paragraphs[i].style.name.startswith('Heading'):
            break
        if doc.paragraphs[i].text.strip():
            count += 1
    print(f'  TABLOLAR LISTESI entries after heading: {count}')
if sl_idx:
    count = 0
    for i in range(sl_idx+1, min(sl_idx+20, len(doc.paragraphs))):
        if doc.paragraphs[i].style.name.startswith('Heading'):
            break
        if doc.paragraphs[i].text.strip():
            count += 1
    print(f'  SEKILLER LISTESI entries after heading: {count}')

print('\n=== SUMMARY ===')
print(f'Total paragraphs: {len(doc.paragraphs)}')
print(f'Total tables: {len(doc.tables)}')
print(f'Total sections: {len(doc.sections)}')
print('DONE')
