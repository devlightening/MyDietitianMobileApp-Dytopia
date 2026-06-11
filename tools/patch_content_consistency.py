from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


DOCX = Path(r"C:\Users\hy971\source\repos\MyDietitianMobileApp\Dytopia_Bitirme_Tezi_TarsusUni_Final_Teknik_Kusursuz.docx")


def normalize_runs(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.italic = False
        run.font.underline = False


doc = Document(DOCX)

replacements = {
    "Haziran 2026, 58 sayfa": "Haziran 2026, 65 sayfa",
    "June 2026, 58 pages": "June 2026, 65 pages",
    (
        "API latency ölçümlerinde P95 değerleri acquisition için 954,7615 ms, normalization için 97,6730 ms, "
        "recommendation için 6,0163 ms ve hybrid recipe için 2,5794 ms olarak ölçülmüştür. Modül envanterinde "
        "50 backend controller, 35 mobil ekran, 35 web panel sayfası ve 45 test dosyası tespit edilmiştir."
    ): (
        "API latency ölçümlerinde P95 değerleri acquisition için 954,7615 ms, normalization için 97,6730 ms, "
        "recommendation için 6,0163 ms ve hybrid recipe için 2,5794 ms olarak ölçülmüştür. Modül envanterinde "
        "50 backend controller, 35 mobil ekran, web panel genelinde 35 sayfa (dashboard altında 24 sayfa) ve "
        "Api.Tests projesinde 33 test/yardımcı C# dosyası tespit edilmiştir."
    ),
}

for paragraph in doc.paragraphs:
    text = paragraph.text
    if text in replacements:
        paragraph.text = replacements[text]
    normalize_runs(paragraph)

for table in doc.tables:
    for row in table.rows:
        if len(row.cells) >= 2 and row.cells[0].text.strip() == "Test dosyası":
            row.cells[1].text = "33"
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                normalize_runs(paragraph)

doc.save(DOCX)
print("content consistency patched")
